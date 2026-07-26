"""
Módulo 09 — Toma de Repertorios / Impresión de Words
Controller: lógica de impresión en segundo plano sin abrir la UI de Word.

Usa Word COM (win32com) para enviar los archivos al spooler de Windows
y win32print para cancelar trabajos pendientes.

Funciones públicas:
  - iniciar_impresion      → lanza el hilo de impresión, retorna total de archivos
  - obtener_estado         → retorna el estado actual (procesados, total, mensaje)
  - cancelar_impresion     → detiene el hilo y elimina los trabajos del spooler
"""

import logging
import threading
from datetime import datetime
from pathlib import Path

import pythoncom
import win32com.client
import win32print
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

import importlib
VBCarpetaConfig = importlib.import_module("modules.02_vb.carpeta_model").VBCarpetaConfig

from .schema import ItemDetalleEnvio

logger = logging.getLogger(__name__)

_CARPETA_MATRIZ = "Vistos Buenos de Abogados"
_NOMBRE_IMPRESORA_DEFECTO = "RICOH IM 550 PCL 6"

# ── Estado global de impresión ────────────────────────────────────────────────
# App de un solo usuario: un diccionario plano protegido por lock es suficiente.
_estado: dict = {
    "en_curso": False,
    "cancelar": False,
    "procesados": 0,
    "total": 0,
    "mensaje": "Listo",
    "nombre_impresora": _NOMBRE_IMPRESORA_DEFECTO,  # se actualiza al iniciar
    "job_id": 0,  # incrementa cada vez que inicia un job nuevo
}
_lock = threading.Lock()


# ── Helpers privados ──────────────────────────────────────────────────────────

def _set_estado(**kwargs) -> None:
    """Actualiza el estado global de impresión de forma segura."""
    with _lock:
        _estado.update(kwargs)


def _obtener_carpeta_dia(db: Session, nombre_carpeta: str | None = None) -> Path:
    """
    Retorna la ruta de la carpeta de Words.
    Si nombre_carpeta se proporciona, usa esa subcarpeta; si no, usa carpeta_hoy del vigilador.
    Lanza ValueError si no está configurada o no existe en disco.
    """
    config = db.query(VBCarpetaConfig).filter(VBCarpetaConfig.id == 1).first()
    if not config or not config.ruta_base:
        raise ValueError(
            "No hay carpeta base configurada. "
            "Activa el vigilador en la secci\u00f3n Carpeta de Descargas VB."
        )
    carpeta = nombre_carpeta or config.carpeta_hoy
    if not carpeta:
        raise ValueError(
            "No hay carpeta del d\u00eda activa. "
            "Selecciona una carpeta o activa el vigilador."
        )
    ruta = Path(config.ruta_base) / _CARPETA_MATRIZ / carpeta
    if not ruta.exists():
        raise ValueError(
            f"La carpeta '{carpeta}' no existe en disco. "
            "Verifica que la carpeta seleccionada sea correcta."
        )
    return ruta


def _leer_nombre_impresora(db: Session) -> str:
    """
    Construye el identificador de impresora a partir de la configuración.
    - Si hay IP y nombre: usa \\\\IP\\nombre (ruta UNC de red)
    - Si solo hay nombre: lo usa directo (impresora instalada en Windows)
    - Si solo hay IP: la usa como identificador
    - Si nada configurado: usa el valor por defecto
    """
    try:
        import importlib
        _mod = importlib.import_module("modules.07_configuracion.model")
        Configuracion = _mod.Configuracion
        cfg = db.query(Configuracion).filter(Configuracion.id == 1).first()
        if cfg:
            ip = (cfg.ip_impresora or "").strip()
            nombre = (cfg.nombre_impresora or "").strip()
            if ip and nombre:
                return f"\\\\{ip}\\{nombre}"
            if nombre:
                return nombre
            if ip:
                return ip
    except Exception as e:
        logger.warning("No se pudo leer configuración de impresora: %s", e)
    return _NOMBRE_IMPRESORA_DEFECTO


def _bucle_impresion(
    archivos: list[Path],
    nombre_impresora: str,
    mi_job_id: int,
    paginas: str | None = None,
) -> None:
    """
    Hilo daemon que imprime los archivos uno a uno usando Word COM.
    Abre una sola instancia de Word (invisible) y la reutiliza para todos.
    Verifica la bandera 'cancelar' entre archivo y archivo.
    Solo actualiza el estado global si sigue siendo el job activo (mi_job_id).
    Si paginas está definido (ej. "1-2"), imprime solo ese rango.
    """
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0   # Suprime diálogos de Word

        # Intentar establecer la impresora objetivo
        try:
            word.ActivePrinter = nombre_impresora
        except Exception as e:
            logger.warning(
                "No se pudo establecer impresora '%s': %s — se usa la predeterminada",
                nombre_impresora, e,
            )

        for i, archivo in enumerate(archivos):
            # Revisar cancelación antes de cada archivo
            with _lock:
                if _estado["cancelar"]:
                    _set_estado(
                        en_curso=False,
                        mensaje=f"Cancelado tras enviar {i} archivo(s) a imprimir",
                    )
                    return

            try:
                doc = word.Documents.Open(str(archivo.resolve()))
                # PrintOut posicional para que COM tome los params correctamente:
                # (Background, Append, Range, OutputFileName, From, To, Item, Copies, Pages)
                # Range=4 → wdPrintRangeOfPages, Item=0 → wdPrintDocumentContent
                if paginas:
                    doc.PrintOut(False, False, 4, "", "", "", 0, 1, paginas)
                else:
                    doc.PrintOut(False)
                doc.Close(SaveChanges=False)
                _set_estado(
                    procesados=i + 1,
                    mensaje=f"Imprimiendo {i + 1} de {len(archivos)}...",
                )
                logger.info("Enviado a imprimir: %s", archivo.name)
            except Exception as e:
                logger.error("Error imprimiendo '%s': %s", archivo.name, e)

    except Exception as e:
        logger.error("Error crítico en hilo de impresión: %s", e)
        _set_estado(mensaje=f"Error: {e}")
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception as e:
            logger.warning("Error al limpiar COM: %s", e)
        # Solo actualizar estado si este hilo sigue siendo el job activo.
        # Evita que un hilo viejo (post-cancelación) sobreescriba el estado
        # de un nuevo job iniciado mientras este hilo aún estaba en finally.
        with _lock:
            if _estado["job_id"] == mi_job_id:
                procesados = _estado["procesados"]
                _set_estado(
                    en_curso=False,
                    mensaje=f"Completado: {procesados} archivo(s) enviados a la impresora",
                )


_MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


# ── Generación del Word de detalle de envío ──────────────────────────────────

def _color_celda(celda, color_hex: str) -> None:
    """Aplica color de fondo (shading) a una celda de tabla."""
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _borde_celda(celda, color_hex: str = "D1D5DB") -> None:
    """Agrega bordes thin en todos los lados de una celda de tabla."""
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        borde = OxmlElement(f"w:{lado}")
        borde.set(qn("w:val"), "single")
        borde.set(qn("w:sz"), "4")
        borde.set(qn("w:color"), color_hex)
        tcBorders.append(borde)
    tcPr.append(tcBorders)


def _parsear_fecha_carpeta(nombre_carpeta: str) -> tuple[int, int, int]:
    """
    Extrae día, mes y año del nombre de carpeta 'VB DD-MM-YYYY'.
    Retorna (dia, mes, anio). Lanza ValueError si no puede parsear.
    """
    fecha_str = nombre_carpeta.replace("VB ", "").strip()
    partes = fecha_str.split("-")
    if len(partes) != 3:
        raise ValueError(f"Formato de fecha no reconocido en '{nombre_carpeta}'")
    return int(partes[0]), int(partes[1]), int(partes[2])


def generar_detalle_envio(
    items: list[ItemDetalleEnvio],
    carpeta_dia: Path,
    nombre_carpeta: str,
) -> str:
    """
    Genera un Word formal con el detalle de los documentos enviados al banco.
    Contiene: encabezado de la notaría, fecha formal, tabla con N°/WF/Cliente y total.
    Se guarda como DETALLE_ENVIO_DD-MM-YYYY.docx en la carpeta del día.
    Retorna el nombre del archivo generado.
    """
    dia, mes, anio = _parsear_fecha_carpeta(nombre_carpeta)
    fecha_formal = f"Santiago, {dia} de {_MESES[mes]} de {anio}"
    fecha_str = nombre_carpeta.replace("VB ", "").strip()

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Encabezado notaría ───────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_before = Pt(0)
    titulo.paragraph_format.space_after = Pt(2)
    run_t = titulo.add_run("33\u00aa Notar\u00eda de Santiago")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(14)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.paragraph_format.space_before = Pt(0)
    subtitulo.paragraph_format.space_after = Pt(4)
    run_s = subtitulo.add_run("Carolina Elizabeth Pi\u00f1a Cuevas \u2013 Notario Titular")
    run_s.font.name = "Calibri"
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # ── Título del documento ─────────────────────────────────────────────────
    titulo_doc = doc.add_paragraph()
    titulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_doc.paragraph_format.space_before = Pt(14)
    titulo_doc.paragraph_format.space_after = Pt(10)
    run_td = titulo_doc.add_run("DETALLE DE ENV\u00cdO DE ESCRITURAS")
    run_td.bold = True
    run_td.font.name = "Calibri"
    run_td.font.size = Pt(13)

    # ── Fecha y destinatario ─────────────────────────────────────────────────
    p_fecha = doc.add_paragraph()
    p_fecha.paragraph_format.space_before = Pt(0)
    p_fecha.paragraph_format.space_after = Pt(10)
    run_f = p_fecha.add_run(fecha_formal)
    run_f.font.name = "Calibri"
    run_f.font.size = Pt(11)

    p_dest = doc.add_paragraph()
    p_dest.paragraph_format.space_before = Pt(0)
    p_dest.paragraph_format.space_after = Pt(0)
    run_d1 = p_dest.add_run("Se\u00f1ores")
    run_d1.font.name = "Calibri"
    run_d1.font.size = Pt(11)

    p_banco = doc.add_paragraph()
    p_banco.paragraph_format.space_before = Pt(0)
    p_banco.paragraph_format.space_after = Pt(0)
    run_d2 = p_banco.add_run("Banco de Chile")
    run_d2.bold = True
    run_d2.font.name = "Calibri"
    run_d2.font.size = Pt(11)

    p_pres = doc.add_paragraph()
    p_pres.paragraph_format.space_before = Pt(0)
    p_pres.paragraph_format.space_after = Pt(10)
    run_d3 = p_pres.add_run("Presente")
    run_d3.font.name = "Calibri"
    run_d3.font.size = Pt(11)
    run_d3.underline = True

    # ── Cuerpo introductorio ─────────────────────────────────────────────────
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(4)
    p_intro.paragraph_format.space_after = Pt(10)
    run_intro = p_intro.add_run(
        "Por medio de la presente, se remiten los siguientes documentos "
        "para su revisi\u00f3n y visaci\u00f3n:"
    )
    run_intro.font.name = "Calibri"
    run_intro.font.size = Pt(11)

    # ── Tabla: N° | WF | Nombre del Cliente ──────────────────────────────────
    anchos = [Inches(0.5), Inches(1.2), Inches(4.8)]
    tabla = doc.add_table(rows=0, cols=3)

    # Encabezado de tabla
    fila_h = tabla.add_row()
    for i, (celda, texto, ancho) in enumerate(
        zip(fila_h.cells, ["N\u00b0", "WF", "Nombre del Cliente"], anchos)
    ):
        celda.width = ancho
        _color_celda(celda, "1565c0")
        _borde_celda(celda)
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(texto)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas de datos
    for idx, item in enumerate(items):
        fondo = "F4F6F8" if idx % 2 == 0 else "FFFFFF"
        fila = tabla.add_row()
        valores = [str(idx + 1), item.wf, item.nombre_cliente]
        alins = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for i, (celda, valor, ancho) in enumerate(zip(fila.cells, valores, anchos)):
            celda.width = ancho
            _color_celda(celda, fondo)
            _borde_celda(celda)
            p = celda.paragraphs[0]
            p.alignment = alins[i]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(valor)
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    # ── Pie con total y timestamp ────────────────────────────────────────────
    total = len(items)
    timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    p_total = doc.add_paragraph()
    p_total.paragraph_format.space_before = Pt(12)
    p_total.paragraph_format.space_after = Pt(0)
    run_total = p_total.add_run(
        f"Total de documentos enviados: {total}"
    )
    run_total.bold = True
    run_total.font.name = "Calibri"
    run_total.font.size = Pt(11)

    p_ts = doc.add_paragraph()
    p_ts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ts.paragraph_format.space_before = Pt(20)
    run_ts = p_ts.add_run(f"Generado: {timestamp}")
    run_ts.font.name = "Calibri"
    run_ts.font.size = Pt(8)
    run_ts.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Guardar ──────────────────────────────────────────────────────────────
    nombre_archivo = f"DETALLE_ENVIO_{fecha_str}.docx"
    ruta = carpeta_dia / nombre_archivo
    try:
        doc.save(str(ruta))
        logger.info("Detalle de envío generado: %s (%d items)", nombre_archivo, total)
    except Exception as e:
        raise ValueError(f"No se pudo generar el detalle de envío: {e}") from e

    return nombre_archivo


# ── Funciones públicas ────────────────────────────────────────────────────────

def iniciar_impresion(
    db: Session,
    paginas: str | None = None,
    items_detalle: list[ItemDetalleEnvio] | None = None,
    nombre_carpeta: str | None = None,
) -> dict:
    """
    Busca los .docx de la carpeta del día (excluye RESUMEN_*, DETALLE_* y subcarpetas)
    y lanza el hilo de impresión.
    Si items_detalle se proporciona, genera un Word de detalle de envío antes de imprimir.
    Si pagina_desde/pagina_hasta están definidos, imprime solo ese rango.
    Lanza ValueError si ya hay una impresión en curso o si no hay archivos.
    Retorna {'total': N, 'mensaje': '...', 'detalle_envio': nombre_archivo | None}.
    """
    with _lock:
        if _estado["en_curso"]:
            raise ValueError("Ya hay una impresión en curso. Espera o cancela primero.")

    carpeta_dia = _obtener_carpeta_dia(db, nombre_carpeta)
    nombre_impresora = _leer_nombre_impresora(db)

    # Excluir RESUMEN, DETALLE_ENVIO y cualquier archivo dentro de subcarpetas
    archivos = sorted([
        f for f in carpeta_dia.glob("*.docx")
        if not f.stem.upper().startswith("RESUMEN")
        and not f.stem.upper().startswith("DETALLE")
    ])

    # Generar Word de detalle de envío si se proporcionaron items (tab Finalización)
    nombre_detalle: str | None = None
    if items_detalle:
        config = db.query(VBCarpetaConfig).filter(VBCarpetaConfig.id == 1).first()
        carpeta_para_detalle = nombre_carpeta or (config.carpeta_hoy if config else "")
        try:
            nombre_detalle = generar_detalle_envio(items_detalle, carpeta_dia, carpeta_para_detalle)
        except Exception as e:
            logger.error("No se pudo generar el detalle de envío: %s", e)

    if not archivos:
        raise ValueError(
            "No hay Words para imprimir en la carpeta del día. "
            "Completa los documentos primero."
        )

    with _lock:
        nuevo_job_id = _estado["job_id"] + 1

    _set_estado(
        en_curso=True,
        cancelar=False,
        procesados=0,
        total=len(archivos),
        mensaje=f"Iniciando impresión de {len(archivos)} archivo(s)...",
        nombre_impresora=nombre_impresora,
        job_id=nuevo_job_id,
    )

    hilo = threading.Thread(
        target=_bucle_impresion,
        args=(archivos, nombre_impresora, nuevo_job_id, paginas),
        daemon=True,
    )
    hilo.start()
    logger.info("Hilo de impresión iniciado: %d archivo(s)", len(archivos))

    return {
        "total": len(archivos),
        "mensaje": _estado["mensaje"],
        "detalle_envio": nombre_detalle,
    }


def obtener_estado() -> dict:
    """Retorna una copia del estado actual de impresión."""
    with _lock:
        return dict(_estado)


def cancelar_impresion() -> dict:
    """
    Detiene el hilo (vía bandera 'cancelar') y elimina los trabajos
    ya enviados al spooler de la impresora.
    Retorna el número de trabajos cancelados en el spooler.
    """
    # Señalar al hilo que debe detenerse en el próximo ciclo
    with _lock:
        _estado["cancelar"] = True

    # Cancelar trabajos que ya llegaron al spooler de Windows
    cancelados_spooler = 0
    with _lock:
        nombre_impresora = _estado.get("nombre_impresora", _NOMBRE_IMPRESORA_DEFECTO)
    try:
        handle = win32print.OpenPrinter(nombre_impresora)
        try:
            jobs = win32print.EnumJobs(handle, 0, -1, 1)
            for job in jobs:
                try:
                    win32print.SetJob(
                        handle, job["JobId"], 0, None, win32print.JOB_CONTROL_DELETE
                    )
                    cancelados_spooler += 1
                except Exception as e:
                    logger.error("No se pudo cancelar job %d: %s", job["JobId"], e)
        finally:
            win32print.ClosePrinter(handle)
    except Exception as e:
        logger.error("Error al acceder al spooler de '%s': %s", nombre_impresora, e)

    mensaje = (
        f"Cancelado. {cancelados_spooler} trabajo(s) eliminados del spooler."
        if cancelados_spooler
        else "Cancelado. No había trabajos en el spooler."
    )
    _set_estado(en_curso=False, mensaje=mensaje)
    logger.info(mensaje)

    return {"cancelados_spooler": cancelados_spooler, "mensaje": mensaje}
