"""
Módulo 09 — Toma de Repertorios
Controller: lógica de negocio pura, sin HTTP.

Funciones públicas:
  - parsear_planilla → lee el Excel 'Consulta OT' del banco, retorna lista de ItemRepertorio
  - llenar_words     → completa los Words de la carpeta del día y genera el Word resumen
                       en la subcarpeta 'Completos'
"""

import io
import logging
import re
import threading
from datetime import date, datetime
from pathlib import Path

import openpyxl
import win32com.client
import win32print
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from num2words import num2words
from sqlalchemy.orm import Session

import importlib
VBCarpetaConfig = importlib.import_module("modules.02_vb.carpeta_model").VBCarpetaConfig

from .schema import ItemRepertorio, ResultadoItem, RespuestaLlenar

logger = logging.getLogger(__name__)

_CARPETA_MATRIZ = "Vistos Buenos de Abogados"
_NOMBRE_COMPLETOS = "Completos"
_NOMBRE_IMPRESORA = "RICOH IM 550 PCL 6"

# ── Estado global de impresión (app de un solo usuario) ───────────────────────
# Persiste entre peticiones HTTP para que el frontend pueda hacer polling.
_estado_impresion: dict = {
    "en_curso": False,
    "cancelar": False,
    "procesados": 0,
    "total": 0,
    "mensaje": "Listo",
}
_lock_impresion = threading.Lock()

_MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


# ── Helpers de conversión de fechas ───────────────────────────────────────────

def _dia_en_letras(dia: int) -> str:
    """Convierte número de día a letras en español. 1 → 'uno'"""
    return num2words(dia, lang="es", to="cardinal")


def _mes_en_texto(mes: int) -> str:
    """Retorna el nombre del mes en español minúsculas."""
    return _MESES.get(mes, str(mes))


# ── Helpers de extracción ─────────────────────────────────────────────────────

def _extraer_wf_archivo(nombre_archivo: str) -> str:
    """
    Extrae el WF del nombre del archivo Word.
    '1044608_20260401094954_abc.docx' → '1044608'
    """
    return Path(nombre_archivo).stem.split("_")[0]


def _extraer_nombre_y_rut(compareciente: str) -> tuple[str, str | None]:
    """
    Extrae nombre y RUT del campo Compareciente del Excel.
    Formato: 'BANCO DE CHILE  97.004.000-5\nNOMBRE CLIENTE  RUT'
    Retorna (nombre_cliente, rut). El RUT puede ser None si no se encuentra.
    """
    if not compareciente:
        return "", None
    lineas = [l.strip() for l in compareciente.replace("\r\n", "\n").split("\n") if l.strip()]
    # La segunda línea tiene formato "NOMBRE CLIENTE  RUT"
    if len(lineas) < 2:
        return "", None
    partes = lineas[1].rsplit("  ", 1)
    if len(partes) > 1:
        return partes[0].strip(), partes[1].strip()
    return lineas[1].strip(), None


# ── Parseo de la planilla Excel del banco ────────────────────────────────────

def parsear_planilla(archivo_bytes: bytes) -> list[ItemRepertorio]:
    """
    Lee el Excel 'Consulta OT' del banco y extrae los datos para llenar los Words.

    Columnas relevantes (base 0):
      2  → Repertorio
      3  → Fecha Repertorio (datetime)
      5  → Cliente (quién encargó: "ROMERO Y ASOCIADOS..." o "BANCO DE CHILE")
      7  → Materia
      8  → Compareciente (BANCO DE CHILE + cliente en líneas separadas)
      22 → Datos Adicionales ('Código de operación: XXXX')

    Ignora encabezados y filas vacías.
    Lanza ValueError si el archivo no puede leerse o no tiene filas válidas.
    """
    try:
        wb = openpyxl.load_workbook(io.BytesIO(archivo_bytes), data_only=True)
    except Exception as e:
        raise ValueError(f"No se pudo abrir el archivo Excel: {e}") from e

    ws = wb.active
    items: list[ItemRepertorio] = []

    for fila in ws.iter_rows(min_row=2, values_only=True):
        if not fila[0]:
            continue

        # WF desde 'Datos Adicionales': "Código de operación: 1044608"
        datos_adicionales = str(fila[22]).strip() if fila[22] else ""
        if ":" not in datos_adicionales:
            continue
        wf = datos_adicionales.split(":")[-1].strip()
        if not wf or not wf.isdigit():
            continue

        # Repertorio y fecha
        repertorio = str(fila[2]).strip() if fila[2] else ""
        fecha_rep = fila[3]
        if not repertorio or not fecha_rep:
            continue

        # Extraer componentes de la fecha
        if isinstance(fecha_rep, (datetime, date)):
            anio = fecha_rep.year
            mes = fecha_rep.month
            dia = fecha_rep.day
        else:
            # Fallback: parsear string si openpyxl no lo convirtió
            try:
                dt = datetime.fromisoformat(str(fecha_rep)[:10])
                anio, mes, dia = dt.year, dt.month, dt.day
            except ValueError:
                logger.warning("No se pudo parsear la fecha '%s' para WF %s", fecha_rep, wf)
                continue

        nombre_cliente, rut = _extraer_nombre_y_rut(str(fila[8]) if fila[8] else "")
        materia = str(fila[7]).strip() if fila[7] else ""

        # Col F: quién encargó el alzamiento — solo se guarda si es un valor conocido
        _cliente_raw = str(fila[5]).strip() if fila[5] else ""
        _CLIENTES_VALIDOS = {"ROMERO Y ASOCIADOS SERVICIOS PROFESIONALES", "BANCO DE CHILE"}
        cliente_notaria = _cliente_raw if _cliente_raw in _CLIENTES_VALIDOS else None

        items.append(ItemRepertorio(
            wf=wf,
            repertorio=repertorio,
            anio=anio,
            dia=dia,
            mes=mes,
            nombre_cliente=nombre_cliente,
            materia=materia,
            rut=rut,
            cliente_notaria=cliente_notaria,
        ))

    if not items:
        raise ValueError(
            "No se encontraron registros válidos en la planilla. "
            "Verifica que sea el formato 'Consulta OT' del banco."
        )

    logger.info("Planilla parseada: %d registros encontrados", len(items))
    return items


# ── Llenado de Words ──────────────────────────────────────────────────────────

def _llenar_parrafo_repertorio(parrafo, repertorio: str, anio: int) -> None:
    """
    Añade el número de repertorio y año al párrafo [0] del Word.
    El texto existe como 'J.E. /N REPERTORIO N°' — se completa con '{rep}-{año}.-'
    Solo actúa si el párrafo aún no fue llenado (no termina en '.-').
    """
    if not parrafo.runs:
        return
    ultimo_run = parrafo.runs[-1]
    texto = ultimo_run.text
    # Detectar si ya está llenado (termina en '.-')
    if texto.rstrip().endswith(".-"):
        return
    # Confirmar que es el párrafo correcto antes de modificar
    if "REPERTORIO" not in texto.upper():
        return
    ultimo_run.text = texto.rstrip() + f"{repertorio}-{anio}.-"


def _llenar_parrafo_fecha(parrafo, dia: int, mes: int) -> None:
    """
    Reemplaza el marcador '____' en el párrafo del cuerpo con la fecha en letras.
    Solo sustituye los guiones, preservando el resto del run (ej: 'del año ...').
    Preserva el formato (bold, fuente, tamaño) del run modificado.
    """
    for run in parrafo.runs:
        if "____" in run.text:
            dia_texto = _dia_en_letras(dia)
            mes_texto = _mes_en_texto(mes)
            run.text = re.sub(r'_{4,}', f"{dia_texto} de {mes_texto}", run.text)
            return


def _llenar_word(ruta_origen: Path, item: ItemRepertorio, ruta_destino: Path) -> None:
    """
    Abre el Word, llena el repertorio en [0] y la fecha en el run con '____'.
    Guarda el resultado en ruta_destino (carpeta Completos).
    Lanza ValueError si no puede abrir o guardar el archivo.
    """
    try:
        doc = Document(str(ruta_origen))
    except Exception as e:
        raise ValueError(f"No se pudo abrir '{ruta_origen.name}': {e}") from e

    # Párrafo [0]: número de repertorio
    if doc.paragraphs:
        _llenar_parrafo_repertorio(doc.paragraphs[0], item.repertorio, item.anio)

    # Buscar la fecha en todos los párrafos (robusto ante variaciones)
    fecha_llenada = False
    for p in doc.paragraphs:
        if any("____" in run.text for run in p.runs):
            _llenar_parrafo_fecha(p, item.dia, item.mes)
            fecha_llenada = True
            break

    if not fecha_llenada:
        logger.warning(
            "No se encontró el marcador '____' en '%s' — fecha no completada",
            ruta_origen.name,
        )

    try:
        doc.save(str(ruta_destino))
    except Exception as e:
        raise ValueError(f"No se pudo guardar '{ruta_destino.name}': {e}") from e


# ── Generación del Word de resumen ────────────────────────────────────────────

def _color_celda(celda, color_hex: str) -> None:
    """Aplica color de fondo (shading) a una celda de tabla."""
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _borde_celda(celda, color_hex: str = "D1D5DB") -> None:
    """Agrega bordes thin en todos los lados de una celda de tabla."""
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        borde = OxmlElement(f"w:{lado}")
        borde.set(qn("w:val"), "single")
        borde.set(qn("w:sz"), "4")
        borde.set(qn("w:color"), color_hex)
        tcBorders.append(borde)
    tcPr.append(tcBorders)


def _agregar_fila_tabla(
    tabla,
    valores: list[str],
    anchos: list,
    es_header: bool = False,
    fondo: str = "FFFFFF",
) -> None:
    """
    Agrega una fila a la tabla del resumen.
    Aplica formato: fuente, tamaño, color de fondo y bordes.
    """
    fila = tabla.add_row()
    alineaciones = [
        WD_ALIGN_PARAGRAPH.CENTER,   # N°
        WD_ALIGN_PARAGRAPH.CENTER,   # WF
        WD_ALIGN_PARAGRAPH.LEFT,     # Nombre Cliente
        WD_ALIGN_PARAGRAPH.CENTER,   # Repertorio
        WD_ALIGN_PARAGRAPH.CENTER,   # Estado
    ]
    for i, (celda, valor, ancho) in enumerate(zip(fila.cells, valores, anchos)):
        celda.width = ancho
        _color_celda(celda, fondo if not es_header else "1565c0")
        _borde_celda(celda)
        parrafo = celda.paragraphs[0]
        parrafo.alignment = alineaciones[i] if i < len(alineaciones) else WD_ALIGN_PARAGRAPH.LEFT
        parrafo.paragraph_format.space_before = Pt(2)
        parrafo.paragraph_format.space_after = Pt(2)
        run = parrafo.add_run(valor)
        run.bold = es_header
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if es_header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _generar_word_resumen(
    resultados: list[ResultadoItem],
    fecha_str: str,
    carpeta_destino: Path,
) -> str:
    """
    Genera el Word de resumen con la misma estructura del PDF de referencia:
      - Título y subtítulo centrados
      - Tabla: N°, WF, Nombre Cliente, Repertorio, Estado
      - Pie con total y timestamp
    Solo incluye los items con estado OK.
    Guarda en carpeta_destino/RESUMEN_{fecha_str}.docx
    Retorna el nombre del archivo generado.
    """
    doc = Document()

    # Márgenes ajustados para que la tabla tenga espacio
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Título ────────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_before = Pt(0)
    titulo.paragraph_format.space_after = Pt(2)
    run_titulo = titulo.add_run("33\u00aa Notar\u00eda de Santiago \u2013 Carolina E. Pi\u00f1a Cuevas")
    run_titulo.bold = True
    run_titulo.font.name = "Calibri"
    run_titulo.font.size = Pt(13)

    # ── Subtítulo ─────────────────────────────────────────────────────────────
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.paragraph_format.space_before = Pt(0)
    subtitulo.paragraph_format.space_after = Pt(10)
    run_sub = subtitulo.add_run(f"Escrituras procesadas VB \u2013 {fecha_str}")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(11)

    # ── Tabla ─────────────────────────────────────────────────────────────────
    # Anchos de columna: N° | WF | Nombre Cliente | Repertorio | Estado
    anchos = [Inches(0.45), Inches(1.05), Inches(3.3), Inches(1.0), Inches(0.7)]
    encabezados = ["N\u00b0", "WF", "Nombre Cliente", "Repertorio", "Estado"]

    tabla = doc.add_table(rows=0, cols=5)

    # Fila de encabezado
    _agregar_fila_tabla(tabla, encabezados, anchos, es_header=True)

    # Filas de datos (solo items OK)
    items_ok = [r for r in resultados if r.estado == "OK"]
    for i, resultado in enumerate(items_ok):
        fondo = "F4F6F8" if i % 2 == 0 else "FFFFFF"
        _agregar_fila_tabla(
            tabla,
            [
                str(i + 1),
                resultado.wf,
                resultado.nombre_cliente,
                resultado.repertorio,
                resultado.estado,
            ],
            anchos,
            fondo=fondo,
        )

    # ── Pie ───────────────────────────────────────────────────────────────────
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before = Pt(10)
    pie.paragraph_format.space_after = Pt(0)
    timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    total = len(items_ok)
    run_pie = pie.add_run(
        f"Total: {total} escritura{'s' if total != 1 else ''} \u2013 Generado: {timestamp}"
    )
    run_pie.font.name = "Calibri"
    run_pie.font.size = Pt(9)
    run_pie.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # ── Guardar ───────────────────────────────────────────────────────────────
    nombre_archivo = f"RESUMEN_{fecha_str}.docx"
    ruta = carpeta_destino / nombre_archivo
    try:
        doc.save(str(ruta))
        logger.info("Word resumen generado: %s (%d registros)", nombre_archivo, total)
    except Exception as e:
        raise ValueError(f"No se pudo generar el Word de resumen: {e}") from e

    return nombre_archivo


# ── Obtener ruta de carpeta activa ────────────────────────────────────────────

def _obtener_ruta_carpeta(db: Session, nombre_carpeta: str | None = None) -> tuple[Path, str]:
    """
    Retorna (ruta_carpeta, fecha_str).
    Si nombre_carpeta se proporciona, busca esa subcarpeta dentro de la carpeta matriz.
    Si no, usa carpeta_hoy del vigilador (comportamiento original).
    Lanza ValueError si no existe o no está configurada.
    """
    config = db.query(VBCarpetaConfig).filter(VBCarpetaConfig.id == 1).first()
    if not config or not config.ruta_base:
        raise ValueError(
            "No hay carpeta base configurada. "
            "Activa el vigilador en la secci\u00f3n Carpeta de Descargas VB."
        )

    carpeta = nombre_carpeta or config.carpeta_hoy
    if not carpeta:
        raise ValueError(
            "No hay carpeta del d\u00eda activa. "
            "Selecciona una carpeta o activa el vigilador."
        )

    ruta = Path(config.ruta_base) / _CARPETA_MATRIZ / carpeta
    if not ruta.exists():
        raise ValueError(
            f"La carpeta '{carpeta}' no existe en disco. "
            "Verifica que la carpeta seleccionada sea correcta."
        )
    # fecha_str desde nombre de carpeta "VB DD-MM-YYYY" → "DD-MM-YYYY"
    fecha_str = carpeta.replace("VB ", "").strip()
    return ruta, fecha_str


# ── Función pública principal ─────────────────────────────────────────────────

def llenar_words(
    items: list[ItemRepertorio],
    db: Session,
    nombre_carpeta: str | None = None,
    anio_override: int | None = None,
) -> RespuestaLlenar:
    """
    Para cada ItemRepertorio, busca el Word en la carpeta indicada (o la del vigilador),
    completa repertorio y fecha, y sobreescribe el archivo original en su lugar.
    Al final genera el Word de resumen en la subcarpeta 'Completos'.

    - nombre_carpeta: subcarpeta VB específica (ej. "VB 14-05-2026"). Si es None usa carpeta_hoy.
    - Los Words sin match en la planilla se ignoran.
    - Los items sin Word quedan como SIN_WORD.
    - Los errores de procesamiento quedan como ERROR.
    Retorna RespuestaLlenar con el detalle completo.
    """
    if not items:
        raise ValueError("No hay items en la planilla para procesar")

    # Determinar el año a usar: override explícito → config guardada → año actual
    if anio_override is None:
        _Configuracion = importlib.import_module("modules.07_configuracion.model").Configuracion
        _config = db.query(_Configuracion).filter(_Configuracion.id == 1).first()
        anio_override = (_config.anio_repertorios if _config and _config.anio_repertorios
                         else datetime.now().year)

    # Reemplazar el año de cada item con el año activo (ignora lo que diga el Excel)
    for item in items:
        item.anio = anio_override

    logger.info("Llenando Words con año fijo: %d", anio_override)

    ruta_carpeta, fecha_str = _obtener_ruta_carpeta(db, nombre_carpeta)

    # Crear subcarpeta Completos solo para el Word de resumen
    carpeta_completos = ruta_carpeta / _NOMBRE_COMPLETOS
    try:
        carpeta_completos.mkdir(exist_ok=True)
    except Exception as e:
        raise ValueError(f"No se pudo crear la carpeta 'Completos': {e}") from e

    # Índice WF → item para búsqueda O(1)
    indice_items: dict[str, ItemRepertorio] = {item.wf: item for item in items}

    # Buscar Words en la carpeta del día (excluye la subcarpeta Completos)
    archivos_word = sorted(
        list(ruta_carpeta.glob("*.docx")) + list(ruta_carpeta.glob("*.doc"))
    )

    resultados: list[ResultadoItem] = []
    wfs_procesados: set[str] = set()

    for archivo in archivos_word:
        wf = _extraer_wf_archivo(archivo.name)
        item = indice_items.get(wf)

        if item is None:
            # Word sin match en la planilla — se deja sin modificar
            logger.info("Word '%s' (WF %s) no tiene match en la planilla, se omite", archivo.name, wf)
            continue

        # Sobreescribe el original en la carpeta del día (reemplaza en su lugar)
        try:
            _llenar_word(archivo, item, archivo)
            resultados.append(ResultadoItem(
                wf=wf,
                nombre_cliente=item.nombre_cliente,
                repertorio=item.repertorio,
                estado="OK",
                nombre_archivo=archivo.name,
            ))
            logger.info("OK: '%s' completado en carpeta del día", archivo.name)
        except Exception as e:
            logger.error("Error procesando '%s': %s", archivo.name, e)
            resultados.append(ResultadoItem(
                wf=wf,
                nombre_cliente=item.nombre_cliente,
                repertorio=item.repertorio,
                estado="ERROR",
                nombre_archivo=archivo.name,
            ))
        finally:
            wfs_procesados.add(wf)

    # Items de la planilla sin Word en la carpeta
    for item in items:
        if item.wf not in wfs_procesados:
            resultados.append(ResultadoItem(
                wf=item.wf,
                nombre_cliente=item.nombre_cliente,
                repertorio=item.repertorio,
                estado="SIN_WORD",
                nombre_archivo=None,
            ))

    total_ok = sum(1 for r in resultados if r.estado == "OK")
    total_sin_word = sum(1 for r in resultados if r.estado == "SIN_WORD")
    total_error = sum(1 for r in resultados if r.estado == "ERROR")

    logger.info(
        "Llenar words completado: %d OK, %d SIN_WORD, %d ERROR → %s",
        total_ok, total_sin_word, total_error, carpeta_completos,
    )

    return RespuestaLlenar(
        resultados=resultados,
        total_ok=total_ok,
        total_sin_word=total_sin_word,
        total_error=total_error,
        ruta_completos=str(carpeta_completos),
        word_resumen="",
    )
