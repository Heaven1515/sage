"""
Módulo 05 — Bóveda
Controller: lógica de negocio pura, sin HTTP.

Funciones públicas:
  - obtener_config         → devuelve o crea el singleton BovedaConfig
  - obtener_entrega_actual → entrega activa de hoy (resetea al cambiar de día)
  - verificar_reingreso    → comprueba si el repertorio ya fue ingresado este año
  - registrar              → crea un nuevo BovedaRegistro
  - obtener_registros      → registros de una fecha concreta
  - editar_registro        → cambia el N° de repertorio (exige clave si es día anterior)
  - eliminar_registro      → borra un registro (exige clave si es día anterior)
  - nueva_entrega          → incrementa el contador de entrega del día
  - buscar_repertorio      → busca en todo el historial
  - generar_word           → genera el Word de la entrega actual, retorna (bytes, nombre)
"""

import io
import logging
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from .model import BovedaConfig, BovedaRegistro
from .schema import RegistroItem

logger = logging.getLogger(__name__)

_CONTRASENA_ADMIN = "notaria33"
_NOTARIA_TITULO   = "33ª Notaría de Santiago – Carolina E. Piña Cuevas"


# ── Helpers de configuración ──────────────────────────────────────────────────

def obtener_config(db: Session) -> BovedaConfig:
    """Retorna el singleton BovedaConfig, creándolo si no existe."""
    config = db.query(BovedaConfig).filter(BovedaConfig.id == 1).first()
    if not config:
        config = BovedaConfig(id=1, entrega_actual=1, fecha_entrega=None)
        db.add(config)
        db.commit()
        db.refresh(config)
    return config


def obtener_entrega_actual(db: Session) -> int:
    """
    Retorna el número de entrega activo para hoy.
    Si la configuración es de un día anterior, la resetea a 1 automáticamente.
    """
    hoy    = date.today()
    config = obtener_config(db)

    if config.fecha_entrega != hoy:
        config.entrega_actual = 1
        config.fecha_entrega  = hoy
        db.commit()

    return config.entrega_actual


# ── Verificación de reingreso ──────────────────────────────────────────────────

def verificar_reingreso(repertorio: int, db: Session) -> dict:
    """
    Verifica si el repertorio ya fue ingresado a bóveda en el año actual.
    Retorna {es_reingreso: bool, registros_previos: list[RegistroItem]}.
    """
    anio = date.today().year
    registros = (
        db.query(BovedaRegistro)
        .filter(
            BovedaRegistro.repertorio == repertorio,
            BovedaRegistro.anio == anio,
        )
        .order_by(BovedaRegistro.fecha_creado)
        .all()
    )
    items = [RegistroItem.model_validate(r) for r in registros]
    return {"es_reingreso": len(items) > 0, "registros_previos": items}


# ── CRUD ───────────────────────────────────────────────────────────────────────

def registrar(repertorio: int, motivo: str | None, db: Session, usuario_id: int | None = None) -> BovedaRegistro:
    """
    Crea un nuevo registro de entrega a bóveda.
    Si el repertorio ya existe en el año actual, requiere un motivo de reingreso.
    Lanza ValueError si el motivo falta en un reingreso o si el repertorio es inválido.
    """
    if repertorio <= 0:
        raise ValueError("El número de repertorio debe ser mayor que cero.")

    hoy  = date.today()
    ahora = datetime.now()

    verificacion = verificar_reingreso(repertorio, db)
    es_reingreso = verificacion["es_reingreso"]

    if es_reingreso and not motivo:
        raise ValueError(
            "El repertorio ya fue ingresado este año. "
            "Debes indicar el motivo del reingreso."
        )

    entrega_actual = obtener_entrega_actual(db)
    motivo_final   = motivo if motivo else "Primer ingreso"

    registro = BovedaRegistro(
        repertorio    = repertorio,
        anio          = hoy.year,
        mes           = hoy.month,
        fecha         = hoy,
        hora_ingreso  = ahora.strftime("%H:%M"),
        es_reingreso  = es_reingreso,
        motivo        = motivo_final,
        numero_entrega = entrega_actual,
        fecha_creado  = ahora,
        usuario_id    = usuario_id,
    )
    db.add(registro)
    db.commit()
    db.refresh(registro)
    logger.info(
        "Bóveda: registrado repertorio %d (reingreso=%s, entrega=%d)",
        repertorio, es_reingreso, entrega_actual,
    )
    return registro


def obtener_registros(fecha: date, db: Session) -> list[BovedaRegistro]:
    """Retorna todos los registros de una fecha, ordenados cronológicamente."""
    return (
        db.query(BovedaRegistro)
        .filter(BovedaRegistro.fecha == fecha)
        .order_by(BovedaRegistro.fecha_creado)
        .all()
    )


def editar_registro(registro_id: int, nuevo_repertorio: int, contrasena: str | None, db: Session) -> BovedaRegistro:
    """
    Cambia el N° de repertorio de un registro existente.
    Exige contraseña de administrador si el registro es de un día anterior.
    Lanza ValueError si no se encuentra, si la contraseña es incorrecta
    o si el nuevo repertorio es inválido.
    """
    registro = db.query(BovedaRegistro).filter(BovedaRegistro.id == registro_id).first()
    if not registro:
        raise ValueError(f"No se encontró el registro con id {registro_id}.")

    if registro.fecha != date.today() and contrasena != _CONTRASENA_ADMIN:
        raise ValueError("Contraseña incorrecta.")

    if nuevo_repertorio <= 0:
        raise ValueError("El número de repertorio debe ser mayor que cero.")

    registro.repertorio = nuevo_repertorio
    db.commit()
    db.refresh(registro)
    logger.info("Bóveda: editado registro %d → repertorio %d", registro_id, nuevo_repertorio)
    return registro


def eliminar_registro(registro_id: int, contrasena: str | None, db: Session) -> None:
    """
    Elimina un registro de bóveda.
    Exige contraseña de administrador si el registro es de un día anterior.
    Lanza ValueError si no se encuentra o si la contraseña es incorrecta.
    """
    registro = db.query(BovedaRegistro).filter(BovedaRegistro.id == registro_id).first()
    if not registro:
        raise ValueError(f"No se encontró el registro con id {registro_id}.")

    if registro.fecha != date.today() and contrasena != _CONTRASENA_ADMIN:
        raise ValueError("Contraseña incorrecta.")

    db.delete(registro)
    db.commit()
    logger.info("Bóveda: eliminado registro %d (repertorio %d)", registro_id, registro.repertorio)


# ── Nueva Entrega ─────────────────────────────────────────────────────────────

def nueva_entrega(db: Session) -> int:
    """
    Incrementa el contador de entregas del día en 1.
    Retorna el nuevo número de entrega activo.
    """
    obtener_entrega_actual(db)  # Sincroniza la fecha antes de incrementar

    config = obtener_config(db)
    config.entrega_actual += 1
    config.fecha_entrega   = date.today()
    db.commit()

    logger.info("Bóveda: nueva entrega activada: %d", config.entrega_actual)
    return config.entrega_actual


# ── Búsqueda ──────────────────────────────────────────────────────────────────

def buscar_repertorio(repertorio: int, db: Session) -> dict:
    """Busca un repertorio en todo el historial de bóveda (sin filtro de año)."""
    registros = (
        db.query(BovedaRegistro)
        .filter(BovedaRegistro.repertorio == repertorio)
        .order_by(BovedaRegistro.fecha_creado)
        .all()
    )
    items = [RegistroItem.model_validate(r) for r in registros]
    return {"repertorio": repertorio, "registros": items, "total": len(items)}


# ── Helpers de formato Word ────────────────────────────────────────────────────

def _color_celda(celda, color_hex: str) -> None:
    """Aplica color de fondo (shading) a una celda de tabla Word."""
    tc   = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _borde_celda(celda) -> None:
    """Agrega bordes finos en todos los lados de una celda."""
    tc        = celda._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        borde = OxmlElement(f"w:{lado}")
        borde.set(qn("w:val"), "single")
        borde.set(qn("w:sz"), "4")
        borde.set(qn("w:color"), "D1D5DB")
        tcBorders.append(borde)
    tcPr.append(tcBorders)


def _agregar_fila_tabla(
    tabla,
    valores: list[str],
    anchos: list,
    es_header: bool = False,
    fondo: str = "FFFFFF",
) -> None:
    """Agrega una fila formateada a la tabla del Word."""
    alineaciones = [
        WD_ALIGN_PARAGRAPH.CENTER,  # N°
        WD_ALIGN_PARAGRAPH.CENTER,  # Repertorio
        WD_ALIGN_PARAGRAPH.CENTER,  # Hora
        WD_ALIGN_PARAGRAPH.LEFT,    # Comentario
    ]
    fila = tabla.add_row()
    for i, (celda, valor, ancho) in enumerate(zip(fila.cells, valores, anchos)):
        celda.width = ancho
        _color_celda(celda, "1565c0" if es_header else fondo)
        _borde_celda(celda)
        parrafo = celda.paragraphs[0]
        parrafo.alignment = alineaciones[i] if i < len(alineaciones) else WD_ALIGN_PARAGRAPH.LEFT
        parrafo.paragraph_format.space_before = Pt(2)
        parrafo.paragraph_format.space_after  = Pt(2)
        run = parrafo.add_run(valor)
        run.bold = es_header
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if es_header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ── Generación del Word ────────────────────────────────────────────────────────

def generar_word(numero_entrega: int, db: Session) -> tuple[bytes, str]:
    """
    Genera el Word de la entrega indicada para el día de hoy.
    Retorna (contenido_bytes, nombre_archivo).
    Lanza ValueError si no hay registros para esa entrega.
    """
    hoy       = date.today()
    registros = (
        db.query(BovedaRegistro)
        .filter(
            BovedaRegistro.fecha == hoy,
            BovedaRegistro.numero_entrega == numero_entrega,
        )
        .order_by(BovedaRegistro.fecha_creado)
        .all()
    )

    if not registros:
        raise ValueError(f"No hay registros para la entrega N° {numero_entrega} del día de hoy.")

    fecha_str = hoy.strftime("%d-%m-%Y")
    doc       = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Título ────────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_before = Pt(0)
    titulo.paragraph_format.space_after  = Pt(2)
    run_titulo = titulo.add_run(_NOTARIA_TITULO)
    run_titulo.bold = True
    run_titulo.font.name = "Calibri"
    run_titulo.font.size = Pt(13)

    # ── Subtítulo ─────────────────────────────────────────────────────────────
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.paragraph_format.space_before = Pt(0)
    subtitulo.paragraph_format.space_after  = Pt(10)
    run_sub = subtitulo.add_run(f"Registro de Bóveda – {fecha_str} – Entrega N° {numero_entrega}")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(11)

    # ── Tabla: N° | Repertorio | Hora | Comentario ────────────────────────────
    anchos      = [Inches(0.45), Inches(1.0), Inches(0.9), Inches(4.15)]
    encabezados = ["N°", "Repertorio", "Hora", "Comentario"]
    tabla = doc.add_table(rows=0, cols=4)

    _agregar_fila_tabla(tabla, encabezados, anchos, es_header=True)

    for i, reg in enumerate(registros):
        fondo = "F4F6F8" if i % 2 == 0 else "FFFFFF"
        _agregar_fila_tabla(
            tabla,
            [str(i + 1), str(reg.repertorio), reg.hora_ingreso, reg.motivo],
            anchos,
            fondo=fondo,
        )

    # ── Pie ───────────────────────────────────────────────────────────────────
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before = Pt(10)
    pie.paragraph_format.space_after  = Pt(0)
    total     = len(registros)
    timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    run_pie   = pie.add_run(
        f"Total: {total} escritura{'s' if total != 1 else ''} – Generado: {timestamp}"
    )
    run_pie.font.name      = "Calibri"
    run_pie.font.size      = Pt(9)
    run_pie.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # ── Guardar en buffer ─────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre = f"BOVEDA_{fecha_str}_ENTREGA_{numero_entrega}.docx"
    logger.info("Word bóveda generado: %s (%d registros)", nombre, total)
    return buffer.read(), nombre
