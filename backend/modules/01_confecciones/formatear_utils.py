"""
Módulo 01 — Confecciones de Borradores
Utilidades de transformación de texto y extracción de datos para el formateo.

Aplica conversiones ortográficas y de forma al texto del borrador sin formato
antes de insertarlo en la escritura pública. Nunca modifica el fondo jurídico:
si dice "fojas 55296", convierte a "fojas cincuenta y cinco mil doscientos
noventa y seis", pero el valor sigue siendo el mismo.

Funciones públicas:
  - transformar_texto  → aplica todas las transformaciones en orden

Funciones privadas de extracción (usadas desde formatear_controller):
  - _extraer_datos_tabla         → borrador del banco (tabla interna)
  - _extraer_titulo              → título del borrador
  - _extraer_clausulas           → cláusulas del borrador
  - _extraer_comuna_conservador  → detecta comuna del Conservador
  - _es_documento_formateado     → detecta si el doc ya fue formateado
  - _extraer_datos_formateado    → datos del encabezado de doc formateado
  - _extraer_clausulas_formateado → cláusulas de doc formateado
"""

import re

from docx import Document
from num2words import num2words


# ── Tabla de meses y dígitos para conversión vehicular y de fechas ───────────

_MESES: dict[str, str] = {
    '01': 'enero',  '02': 'febrero', '03': 'marzo',     '04': 'abril',
    '05': 'mayo',   '06': 'junio',   '07': 'julio',     '08': 'agosto',
    '09': 'septiembre', '10': 'octubre', '11': 'noviembre', '12': 'diciembre',
}

_DIGITOS: dict[str, str] = {
    '0': 'cero', '1': 'uno',   '2': 'dos',   '3': 'tres', '4': 'cuatro',
    '5': 'cinco','6': 'seis',  '7': 'siete', '8': 'ocho', '9': 'nueve',
}

# ── Tabla de abreviaturas notariales ─────────────────────────────────────────
# Pares (patron_regex, reemplazo). Orden importante: más específicos primero.
_ABREVIATURAS: list[tuple[str, str]] = [
    # Vuelta (reverso de página en registros)
    (r'\bVTA\b',  'Vuelta'),
    (r'\bVta\b',  'Vuelta'),
    (r'\bvta\.',  'vuelta'),
    (r'\bvta\b',  'vuelta'),
    # Porcentaje
    (r'%',        'porciento'),
    # Número — todas las variantes tipográficas
    (r'N°',       'número'),
    (r'n°',       'número'),
    (r'Nº',       'número'),
    (r'nº',       'número'),
    # Artículo — número de artículo en registros
    (r'\bart\.',  'artículo'),
    (r'\bArt\.',  'Artículo'),
    # Fojas — abreviatura registral
    (r'\bfs\.',   'fojas'),
    (r'\bFs\.',   'Fojas'),
]


# ── Equiparación de tildes para búsqueda de nombres ──────────────────────────

def _patron_ignorar_tildes(texto: str) -> str:
    """
    Convierte cada letra a una clase de caracteres que equipara la versión
    sin tilde y con tilde. Permite que "HERNANDEZ" matchee "Hernández".
    """
    _MAPA = {
        'a': '[aáàâäAÁÀÂÄ]', 'e': '[eéèêëEÉÈÊË]',
        'i': '[iíìîïIÍÌÎÏ]', 'o': '[oóòôöOÓÒÔÖ]',
        'u': '[uúùûüUÚÙÛÜ]', 'n': '[nñNÑ]',
    }
    resultado = []
    for c in texto:
        cl = c.lower()
        if cl in _MAPA:
            resultado.append(_MAPA[cl])
        elif c.isalpha():
            resultado.append(f'[{c.upper()}{c.lower()}]')
        else:
            resultado.append(re.escape(c))
    return ''.join(resultado)


def _patron_nombre_flexible(nombre: str) -> str:
    """
    Patrón regex que localiza el nombre del banco en el cuerpo del borrador.
    - Tolera diferencias de tilde (MUNOZ ≡ MUÑOZ).
    - Para nombres de 3 palabras (NOMBRE APELLIDO1 APELLIDO2), permite un
      segundo nombre opcional entre el primer nombre y los apellidos,
      cubriendo el caso en que el banco guarda la forma corta pero el
      borrador incluye el nombre completo (ej: "Gabriel Enrique Romero Silva").
    - Para nombres de 4+ palabras (personas con 4 componentes o razones sociales),
      une las palabras con \\s+ (maneja dobles espacios del borrador) y permite
      una 's' final opcional por palabra para cubrir variantes singular/plural
      (ej: ASESORIA ↔ ASESORIAS, ASESORIA E INVERSIONES WIBO LIMITADA).
    """
    partes = nombre.strip().upper().split()
    if len(partes) == 3:
        return (
            _patron_ignorar_tildes(partes[0]) +
            r'(?:\s+[A-ZÁÉÍÓÚÑa-záéíóúñ]+)?\s+' +
            _patron_ignorar_tildes(partes[1]) + r'\s+' +
            _patron_ignorar_tildes(partes[2])
        )
    if len(partes) >= 4:
        return r'\s+'.join(_patron_ignorar_tildes(p) + r'[sS]?' for p in partes)
    return _patron_ignorar_tildes(nombre)


# ── Conversión de fechas y códigos vehiculares ────────────────────────────────

def _codigo_alfanumerico(codigo: str) -> str:
    """
    Convierte un código alfanumérico para escritura pública:
    - Grupos de letras consecutivas → se mantienen juntas (ej: KJT → KJT)
    - Dígitos → se escriben uno a uno (ej: 10 → uno cero)
    - Punto → 'punto'
    - Guion → 'guion'
    Ejemplo: JK100X → JK uno cero cero X
    """
    tokens = re.findall(r'[A-Za-záéíóúñÁÉÍÓÚÑ]+|\d|[.\-]', codigo)
    resultado: list[str] = []
    for token in tokens:
        if token in _DIGITOS:
            resultado.append(_DIGITOS[token])
        elif token == '.':
            resultado.append('punto')
        elif token == '-':
            resultado.append('guion')
        else:
            resultado.append(token.upper())
    return ' '.join(resultado)


def _convertir_ruts(texto: str) -> str:
    """
    Convierte RUT chileno (XX.XXX.XXX-X) a su forma escrita completa.
    Debe ejecutarse ANTES de _convertir_numeros para que los puntos del RUT
    no hagan que cada segmento se convierta por separado.
    Ejemplo: 13.499.632-3 → trece millones cuatrocientos noventa y nueve mil
                             seiscientos treinta y dos guion tres
    El dígito verificador K se deja como letra.
    """
    def _reemplazar(m: re.Match) -> str:
        numero      = int(m.group(1).replace('.', ''))
        verificador = m.group(2).upper()
        en_letras   = num2words(numero, lang='es')
        ver_letras  = 'K' if verificador == 'K' else num2words(int(verificador), lang='es')
        return f'{en_letras} guion {ver_letras}'

    return re.sub(r'(\d{1,2}\.\d{3}\.\d{3})-([0-9kK])', _reemplazar, texto)


def _convertir_fechas_numericas(texto: str) -> str:
    """
    Convierte fechas en formato DD-MM-YYYY a su escritura en letras.
    Ejemplo: 14-07-2009 → catorce de julio del año dos mil nueve
    Se aplica antes de _convertir_numeros para evitar conversión parcial.
    """
    def _reemplazar(m: re.Match) -> str:
        dia  = num2words(int(m.group(1)), lang='es')
        mes  = _MESES.get(m.group(2).zfill(2), m.group(2))
        anio = num2words(int(m.group(3)), lang='es')
        return f'{dia} de {mes} del año {anio}'

    return re.sub(r'\b(\d{1,2})-(\d{2})-(\d{4})\b', _reemplazar, texto)


def _convertir_datos_vehiculo(texto: str) -> str:
    """
    Convierte dígito a dígito los campos de identificación de vehículos:
    número de motor, número de chasis, modelo y patente/inscripción.
    Debe ejecutarse ANTES de _convertir_numeros.
    """
    # Número de motor: código sin espacios (letras + dígitos)
    texto = re.sub(
        r'(número de motor\s+)([A-Z0-9]+)',
        lambda m: m.group(1) + _codigo_alfanumerico(m.group(2)),
        texto, flags=re.IGNORECASE,
    )
    # Número de chasis: mismo formato que motor
    texto = re.sub(
        r'(número de chasis\s+)([A-Z0-9]+)',
        lambda m: m.group(1) + _codigo_alfanumerico(m.group(2)),
        texto, flags=re.IGNORECASE,
    )
    # Patente: código con posibles puntos y guiones (ej: UK.2696-4)
    texto = re.sub(
        r'(patente o inscripción\s+)([A-Z0-9.\-]+)',
        lambda m: m.group(1) + _codigo_alfanumerico(m.group(2)),
        texto, flags=re.IGNORECASE,
    )
    # Modelo: solo activa si el valor contiene al menos un dígito (ej: 496 ST 1.8).
    # Esto evita falsos positivos en documentos que usen "modelo" en contexto textual.
    def _reemplazar_modelo(m: re.Match) -> str:
        valor = m.group(2).strip()
        if not any(c.isdigit() for c in valor):
            return m.group(0)  # no tiene dígitos → no es código vehicular
        return m.group(1) + ' '.join(_codigo_alfanumerico(s) for s in valor.split())

    texto = re.sub(
        r'(modelo\s+)([^,]+?)(?=,)',
        _reemplazar_modelo,
        texto, flags=re.IGNORECASE,
    )
    return texto


# ── Conversión de números a letras ────────────────────────────────────────────

def _convertir_abreviaturas(texto: str) -> str:
    """Reemplaza abreviaturas notariales por su forma escrita completa."""
    for patron, reemplazo in _ABREVIATURAS:
        texto = re.sub(patron, reemplazo, texto)
    return texto



def _convertir_numeros(texto: str) -> str:
    """
    Reemplaza todos los números arábigos por su equivalente en español.

    Casos manejados:
    - Enteros simples:           55296 → cincuenta y cinco mil doscientos noventa y seis
    - Miles con punto (chileno): 27.000 → veintisiete mil
    - Miles + decimal con coma:  27.000,50 → veintisiete mil coma cincuenta
    - Decimales con coma:        493,30 → cuatrocientos noventa y tres coma treinta
    - Números en códigos:        B-03 → B-tres  (solo se convierte el dígito)

    Los textos ya escritos en letras no se tocan.
    """
    # Miles con punto (separador chileno), con decimal opcional: 27.000 o 27.000,50
    # Debe ir antes del decimal-coma para que "27.000,50" no se parta en "000,50"
    def _reemplazar_miles(m: re.Match) -> str:
        texto_num = m.group()
        if ',' in texto_num:
            partes        = texto_num.split(',')
            parte_entera  = int(partes[0].replace('.', ''))
            parte_decimal = partes[1]
            return (
                f"{num2words(parte_entera, lang='es')} coma "
                f"{num2words(int(parte_decimal), lang='es')}"
            )
        return num2words(int(texto_num.replace('.', '')), lang='es')

    texto = re.sub(r'\b\d{1,3}(?:\.\d{3})+(?:,\d+)?\b', _reemplazar_miles, texto)

    # Decimales con coma deben procesarse antes que los enteros (ej: "493,30")
    def _reemplazar_decimal(m: re.Match) -> str:
        parte_entera  = int(m.group(1))
        parte_decimal = m.group(2)
        return (
            f"{num2words(parte_entera, lang='es')} coma "
            f"{num2words(int(parte_decimal), lang='es')}"
        )

    texto = re.sub(r'(\d+),(\d+)', _reemplazar_decimal, texto)

    # Enteros — \b garantiza que no estén dentro de palabras
    texto = re.sub(
        r'\b\d+\b',
        lambda m: num2words(int(m.group()), lang='es'),
        texto,
    )
    return texto


# ── Corrección tipográfica ────────────────────────────────────────────────────

def _limpiar_puntuacion(texto: str) -> str:
    """
    Corrige problemas tipográficos frecuentes en borradores del banco:
    - Punto o coma sin espacio posterior → agrega espacio
    - Espacios múltiples consecutivos   → un solo espacio
    - Espacio justo antes de puntuación → elimina
    """
    # Normalizar espacios de no separación (U+00A0) y otros Unicode a espacio normal.
    # Word usa U+00A0 para alinear texto; sobreviven a la extracción y rompen los
    # patrones de búsqueda que solo contemplan el espacio ASCII (U+0020).
    texto = re.sub(r'[^\S\n\r]', ' ', texto)
    # Punto/coma/punto y coma sin espacio a la derecha (excepto ante otro signo o dígito)
    texto = re.sub(r'([.,;:])(?=[^\s.,;:\d\-])', r'\1 ', texto)
    # Espacios dobles o más
    texto = re.sub(r' {2,}', ' ', texto)
    # Espacio antes de puntuación
    texto = re.sub(r' ([.,;:])', r'\1', texto)
    return texto.strip()


# ── Función pública de transformación ────────────────────────────────────────

def transformar_texto(texto: str) -> str:
    """
    Aplica todas las transformaciones de forma al texto del borrador.

    Orden:
    1. Limpieza de puntuación (para que las abreviaturas no tengan basura)
    2. Expansión de abreviaturas notariales
    3. Conversión de fechas DD-MM-YYYY a letras (antes que _convertir_numeros)
    4. Conversión de códigos vehiculares dígito a dígito (antes que _convertir_numeros)
    5. Conversión de números a letras
    6. Segunda limpieza (la conversión puede generar espacios extra)

    No altera los valores jurídicos: fojas, números de inscripción, años, etc.
    solo cambian de dígitos a palabras, pero el contenido es idéntico.
    """
    texto = _limpiar_puntuacion(texto)
    texto = _convertir_abreviaturas(texto)
    texto = _convertir_ruts(texto)
    texto = _convertir_fechas_numericas(texto)
    texto = _convertir_datos_vehiculo(texto)
    texto = _convertir_numeros(texto)
    texto = _limpiar_puntuacion(texto)
    return texto


# ── Extracción desde borrador del banco (sin formato) ────────────────────────

def _extraer_datos_tabla(doc: Document) -> dict:
    """
    Lee la tabla interna del borrador sin formato y extrae los datos clave.
    La tabla tiene el formato:
      Fila 0: WF            | <número>
      Fila 1: Abogado Red.  | <nombre>
      Fila 2: Cliente       | <APELLIDO NOMBRE>
      Fila 3: Rut cliente   | <RUT>
    """
    if not doc.tables:
        raise ValueError("El documento no contiene la tabla de datos esperada.")

    tabla = doc.tables[0]
    datos: dict[str, str] = {}

    for fila in tabla.rows:
        if len(fila.cells) < 2:
            continue
        clave = fila.cells[0].text.strip().lower()
        valor = fila.cells[1].text.strip()
        if "wf" in clave or "carpeta" in clave:
            datos["wf"] = valor
        elif "abogado" in clave:
            datos["abogado_redactor"] = valor
        elif "cliente" in clave and "rut" not in clave:
            datos["nombre_cliente"] = valor
        elif "rut" in clave:
            datos["rut_cliente"] = valor

    for campo in ["wf", "abogado_redactor", "nombre_cliente"]:
        if campo not in datos:
            raise ValueError(f"No se encontró el campo '{campo}' en la tabla del borrador.")

    return datos


def _extraer_titulo(doc: Document) -> str:
    """
    Extrae el título de la escritura — el primer párrafo no vacío antes
    de la comparecencia. Respeta cualquier tipo de escritura tal como
    viene del banco (alzamiento de hipoteca, cancelación, prenda, etc.).
    Retorna cadena vacía si no se encuentra — nunca impone un título por defecto.
    """
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto or texto == "\t":
            continue
        if texto.upper().startswith("COMPARECE"):
            break
        return texto.upper()
    return ""


def _extraer_clausulas(doc: Document) -> str:
    """
    Reúne los párrafos de cláusulas del borrador en un único string.
    Omite: tabla, párrafos vacíos/tabs, título y párrafo de comparecencia.
    Los párrafos se unen con espacio — en la escritura pública van como
    un solo párrafo continuo.

    El título se identifica como el primer párrafo no vacío antes de la
    comparecencia, sin importar su contenido (alzamiento, cancelación, etc.).

    Normalización de mayúsculas: el banco manda descripciones de inmuebles
    en ALL CAPS. Si un párrafo tiene >80 % de sus letras en mayúscula y no
    es una etiqueta de cláusula (PRIMERO/SEGUNDO/TERCERO), se convierte a
    título para mantener el cuerpo visualmente uniforme.
    """
    ETIQUETAS_CLAUSULA = ("PRIMERO:", "SEGUNDO:", "TERCERO:", "CUARTO:", "QUINTO:", "SEXTO:")
    ignorar = {"", "\t"}
    textos: list[str] = []
    titulo_omitido = False  # El título es el primer párrafo no vacío antes de la comparecencia

    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto in ignorar:
            continue
        # Omitir el título (sea cual sea su contenido)
        if not titulo_omitido and not texto.upper().startswith("COMPARECE"):
            titulo_omitido = True
            continue
        if texto.upper().startswith("COMPARECE"):
            continue
        if texto.upper().startswith("PERSONERÍAS") or texto.upper().startswith("PERSONERIAS"):
            continue
        # Detectar y normalizar párrafos ALL CAPS del banco (ej: descripción del inmueble)
        letras = [c for c in texto if c.isalpha()]
        if letras and not texto.upper().startswith(ETIQUETAS_CLAUSULA):
            ratio_caps = sum(1 for c in letras if c.isupper()) / len(letras)
            if ratio_caps > 0.8:
                texto = texto.title()
        textos.append(texto)

    return " ".join(textos)


def _extraer_comuna_conservador(doc: Document) -> list[str]:
    """
    Detecta todas las comunas del Conservador de Bienes Raíces en el documento.
    Retorna lista ordenada y deduplicada (preserva orden de aparición).
    Cubre documentos con uno o varios conservadores (ej: prenda con 2 comunas).
    """
    patron = r'Conservador de Bienes Ra[ií]ces de\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ ]+?)(?=[,.]|$)'

    def _buscar_todos(texto: str) -> list[str]:
        return [m.group(1).strip().upper()
                for m in re.finditer(patron, texto.strip(), re.UNICODE | re.IGNORECASE)]

    vistas: list[str] = []
    seen:   set[str]  = set()

    def _agregar(comunas: list[str]) -> None:
        for c in comunas:
            if c not in seen:
                seen.add(c)
                vistas.append(c)

    for p in doc.paragraphs:
        _agregar(_buscar_todos(p.text))

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    _agregar(_buscar_todos(p.text))

    return vistas


# ── Extracción desde documento ya formateado ─────────────────────────────────

def _es_documento_formateado(doc: Document) -> bool:
    """Detecta si el doc ya fue formateado por este sistema: no tiene tabla interna."""
    return not doc.tables


def _extraer_datos_formateado(doc: Document) -> dict:
    """
    Extrae los datos del encabezado de un documento ya formateado.
    Estructura de párrafos del cuerpo (los no vacíos, en orden):
      0: J.E. /6 REPERTORIO Nº
      1: Abogado Redactor: <nombre>
      2: WF <número>
      3: <COMUNA>
      4: <TITULO>
      5: BANCO DE CHILE
      6: A
      7: <NOMBRE_CLIENTE>
      8: <cuerpo largo>  ...
    Retorna las mismas claves que _extraer_datos_tabla, más 'titulo' y 'comuna'.
    """
    textos = [p.text.strip() for p in doc.paragraphs]
    datos: dict[str, str] = {}

    for texto in textos:
        if texto.startswith("Abogado Redactor:"):
            datos["abogado_redactor"] = texto.split(":", 1)[1].strip()
        elif texto.startswith("WF "):
            datos["wf"] = texto[3:].strip()

    # Nombre cliente: párrafo inmediatamente después del párrafo "A"
    for i, texto in enumerate(textos):
        if texto == "A" and i + 1 < len(textos):
            datos["nombre_cliente"] = textos[i + 1]
            break

    # Extracción posicional sobre párrafos no vacíos (estructura del doc formateado):
    #   Formato actual: 0: REPERTORIO  1: Abogado  2: WF  3: COMUNA  4: TITULO  5: BANCO DE CHILE
    #   Formato viejo:  0: REPERTORIO  1: Abogado  2: WF  3: TITULO  4: BANCO DE CHILE
    # Detección confiable: si pos4 es "BANCO DE CHILE", entonces pos3 es el título (formato viejo).
    # El enfoque anterior usaba palabras clave en pos3, pero "PRENDA" es también identificador
    # de tipo de documento (no título) y generaba falsos positivos.
    no_vacios = [t for t in textos if t]
    pos4 = no_vacios[4].upper() if len(no_vacios) > 4 else ""
    if pos4.startswith("BANCO DE CHILE"):
        # Formato viejo: posición 3 es el título, no hay línea de comuna
        datos["titulo"] = no_vacios[3] if len(no_vacios) > 3 else ""
        datos["comuna"] = ""
    else:
        # Formato actual: posición 3 es la comuna/identificador, posición 4 es el título
        datos["comuna"] = no_vacios[3] if len(no_vacios) > 3 else ""
        datos.setdefault("titulo", no_vacios[4] if len(no_vacios) > 4 else "")

    for campo in ["wf", "abogado_redactor", "nombre_cliente"]:
        if campo not in datos:
            raise ValueError(f"No se encontró '{campo}' en el documento formateado.")
    return datos


def _extraer_clausulas_formateado(doc: Document) -> str:
    """
    Extrae las cláusulas del párrafo de cuerpo de un documento ya formateado.
    Las cláusulas están entre 'exponen: ' y la primera aparición de 'LA PERSONERÍA'.
    El texto ya está transformado (números en letras), no requiere conversión.
    """
    for p in doc.paragraphs:
        texto = p.text
        # Acepta tanto "exponen: " (varios comparecientes) como "expone: " (uno solo)
        m = re.search(r'exponen?\s*:\s*', texto, re.IGNORECASE)
        if not m:
            continue
        resto = texto[m.end():]
        fin1 = resto.upper().find("LA PERSONERÍA")
        fin2 = resto.upper().find("PERSONERÍAS:")
        candidatos = [f for f in [fin1, fin2] if f != -1]
        fin = min(candidatos) if candidatos else -1
        return resto[:fin].strip() if fin != -1 else resto.strip()
    raise ValueError("No se encontró el cuerpo de cláusulas en el documento formateado.")


# ── Detección de ordinales para cláusula automática ──────────────────────────

_UNIDADES_ORD: dict[str, int] = {
    'PRIMERO': 1, 'SEGUNDO': 2, 'TERCERO': 3, 'CUARTO': 4,
    'QUINTO': 5, 'SEXTO': 6, 'SEPTIMO': 7, 'OCTAVO': 8, 'NOVENO': 9,
}
_DECENAS_ORD: dict[str, int] = {
    'DECIMO': 10, 'VIGESIMO': 20, 'TRIGESIMO': 30, 'CUADRAGESIMO': 40,
    'QUINCUAGESIMO': 50, 'SEXAGESIMO': 60, 'SEPTUAGESIMO': 70,
    'OCTAGESIMO': 80, 'NONAGESIMO': 90, 'CENTESIMO': 100,
}
_UNIDADES_INV: dict[int, str] = {v: k for k, v in _UNIDADES_ORD.items()}
_DECENAS_INV: dict[int, str] = {v: k for k, v in _DECENAS_ORD.items()}


def _ordinal_a_numero(texto: str) -> int:
    """Convierte texto ordinal en español a número (ej: 'DECIMO TERCERO' → 13)."""
    total = 0
    for parte in texto.strip().upper().replace('SÉPTIMO', 'SEPTIMO').split():
        total += _DECENAS_ORD.get(parte, 0) + _UNIDADES_ORD.get(parte, 0)
    return total


def _numero_a_ordinal(n: int) -> str:
    """Convierte número a ordinal en español mayúsculas (ej: 126 → 'CENTESIMO VIGESIMO SEXTO')."""
    if n <= 0:
        return 'PRIMERO'
    partes: list[str] = []
    if n >= 100:
        partes.append('CENTESIMO')
        n -= 100
    decenas = (n // 10) * 10
    if decenas and decenas in _DECENAS_INV:
        partes.append(_DECENAS_INV[decenas])
        n -= decenas
    if n and n in _UNIDADES_INV:
        partes.append(_UNIDADES_INV[n])
    return ' '.join(partes) if partes else 'PRIMERO'


def _detectar_siguiente_clausula(clausulas_tx: str) -> str:
    """
    Encuentra el ordinal más alto en el texto y retorna el siguiente.
    Usa el MÁXIMO encontrado (no el último), para manejar borradores que
    repiten un ordinal por error (ej: dos cláusulas TERCERO consecutivas).
    Usado para numerar automáticamente la cláusula del art. 401 N°12 COT.
    """
    _DEC = (
        r'(?:CENTESIMO|NONAGESIMO|OCTAGESIMO|SEPTUAGESIMO|SEXAGESIMO|'
        r'QUINCUAGESIMO|CUADRAGESIMO|TRIGESIMO|VIGESIMO|DECIMO)'
    )
    _UNI = r'(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|S[EÉ]PTIMO|OCTAVO|NOVENO)'
    patron = rf'({_DEC}(?:\s+{_DEC})?(?:\s+{_UNI})?|{_UNI}):'

    matches = list(re.finditer(patron, clausulas_tx))
    if not matches:
        return 'SEGUNDO'

    # Usar el ordinal más alto para evitar error cuando el borrador repite un ordinal
    numeros = [
        _ordinal_a_numero(m.group(1).strip().replace('SÉPTIMO', 'SEPTIMO'))
        for m in matches
    ]
    maximo = max(numeros)
    return _numero_a_ordinal(maximo + 1)
