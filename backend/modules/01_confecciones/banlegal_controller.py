"""
Módulo 01 — Confecciones de Borradores
Controller de formateo para borradores Banlegal (alzamientos de prenda sobre vehículo).

Los borradores Banlegal difieren del formato estándar BDC:
  - Sin tabla interna: datos en párrafos (ABOGADO:, WF:)
  - Un solo apoderado propio (no los combos Nathalie/Jonatan/Andrea)
  - Cliente puede ser persona jurídica (empresa), sin inversión de nombre
  - Encabezado muestra 'BANLEGAL' en lugar de la comuna

Función pública:
  - formatear_banlegal(contenido, apo_id, notario) → bytes
"""

import io
import logging
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .formatear_utils import (
    transformar_texto,
    _extraer_comuna_conservador,
)

logger = logging.getLogger(__name__)


# ── Apoderados Banlegal ───────────────────────────────────────────────────────
# Fuente: escritura pública de personería del 18 de abril de 2022,
# Notaría de Santiago de don Gerardo Carvallo Castillo.
_APODERADOS_BANLEGAL: dict[str, dict] = {
    "ronnie": {
        "nombre":             "RONNIE ESTEBAN PALMA CARTES",
        "tratamiento":        "don",
        "descripcion":        (
            "chileno, soltero, abogado, cédula de identidad número doce millones "
            "quinientos cuatro mil cincuenta y nueve guion cuatro"
        ),
        "fecha_personeria":   "dieciocho de abril del año dos mil veintidós",
        "notaria_personeria": "la Notaría de Santiago de don Gerardo Carvallo Castillo",
    },
    "felix": {
        "nombre":             "FÉLIX FRANCISCO MORALES TRUJILLO",
        "tratamiento":        "don",
        "descripcion":        (
            "chileno, soltero, abogado, cédula de identidad doce millones cuatrocientos "
            "cuarenta y cuatro mil cuatrocientos setenta y nueve guion nueve"
        ),
        "fecha_personeria":   "dieciocho de abril del año dos mil veintidós",
        "notaria_personeria": "la Notaría de Santiago de don Gerardo Carvallo Castillo",
    },
}


# ── Helpers de formato (réplica mínima para no crear dependencia circular) ────

def _fuente(run) -> None:
    run.font.name = "Courier New"
    run.font.size = Pt(12.5)
    # Fijar idioma español de Chile para que Word no subraye el texto en rojo
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'es-CL')
    rPr.append(lang)


def _run(parrafo, texto: str, negrita=None, subrayado=None) -> None:
    if not texto:
        return
    r = parrafo.add_run(texto)
    _fuente(r)
    r.bold = negrita
    if subrayado:
        r.underline = subrayado


def _parrafo(doc, texto: str = "", negrita=None,
             alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY, interlineado: float = 2.0):
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.line_spacing = interlineado
    p.paragraph_format.space_after  = Pt(0)
    if texto:
        _run(p, texto, negrita=negrita)
    return p


# ── Detección y extracción específica para banlegal ──────────────────────────

# Líneas de cabecera que preceden al título real en los borradores banlegal.
# Aplica tanto al formato "BANLEGAL" (prenda) como al formato "Fiscalía" (hipoteca).
_CABECERA_PREFIJOS = ("REPERTORIO", "ABOGADO", "WF", "BANLEGAL", "BANCO", "FISCAL")


def _es_banlegal_formateado(doc: Document) -> bool:
    """
    Detecta si el doc ya fue formateado por SAGE para banlegal.
    Los formateados tienen 'Abogado Redactor:' en el encabezado;
    los borradores crudos tienen 'ABOGADO:' (mayúsculas, sin 'Redactor').
    No se puede usar 'not doc.tables' porque los borradores banlegal tampoco
    tienen tabla interna, a diferencia de los borradores BDC estándar.
    """
    return any(
        p.text.strip().startswith("Abogado Redactor:")
        for p in doc.paragraphs
    )


def _extraer_titulo_banlegal(doc: Document) -> str:
    """
    Extrae el título real del borrador banlegal saltando las líneas de cabecera.
    El título es el primer párrafo no vacío que no sea cabecera ni comparecencia.
    Salta: REPERTORIO, ABOGADO, WF, BANLEGAL, BANCO (DE CHILE), Fiscalía, etc.
    """
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto or texto.startswith("_"):
            continue
        tu = texto.upper()
        if tu == "A":
            continue
        if tu.startswith("COMPARECE"):
            break
        if any(tu.startswith(pref) for pref in _CABECERA_PREFIJOS):
            continue
        return tu
    return ""


def _extraer_datos_banlegal(doc: Document) -> dict:
    """
    Extrae WF y abogado de borradores banlegal (sin tabla interna).
    Acepta 'WF: XXXX' y 'WF XXXX' (con o sin dos puntos).
    """
    datos: dict[str, str] = {}
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto.upper().startswith("ABOGADO:"):
            datos["abogado_redactor"] = texto.split(":", 1)[1].strip()
        elif re.match(r'^WF[\s:]+', texto, re.IGNORECASE):
            # Elimina el prefijo 'WF' seguido de ':' o espacio(s)
            datos["wf"] = re.sub(r'^WF[\s:]+', '', texto, flags=re.IGNORECASE).strip()
    for campo in ["wf", "abogado_redactor"]:
        if campo not in datos:
            raise ValueError(f"No se encontró '{campo}' en el borrador banlegal.")
    return datos


def _extraer_clausulas_banlegal(doc: Document) -> str:
    """
    Extrae cláusulas de borrador banlegal.
    Reúne párrafos desde el primer 'PRIMERO:' hasta 'LA PERSONERÍA'.
    """
    en_clausulas = False
    textos: list[str] = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue
        if texto.upper().startswith("PRIMERO"):
            en_clausulas = True
        if texto.upper().startswith("LA PERSON"):
            break
        if en_clausulas:
            textos.append(texto)
    return " ".join(textos)


def _extraer_cliente_banlegal(clausulas: str) -> str:
    """
    Intenta extraer el nombre del cliente de las cláusulas.
    Busca el patrón 'la sociedad NOMBRE, rol único'.
    Retorna '_________________________' si no encuentra.
    """
    m = re.search(
        r'la sociedad\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+?)(?=,\s*rol)',
        clausulas, re.IGNORECASE,
    )
    if m:
        return m.group(1).strip().upper()
    return "_________________________"


# ── Construcción del documento ────────────────────────────────────────────────

def _comparecencia_banlegal(parrafo, notario: dict, apo: dict, anio: str) -> None:
    """Comparecencia con un solo apoderado banlegal."""
    _run(parrafo, "EN SANTIAGO DE CHILE, ", negrita=True)
    _run(parrafo, f"a ____________________ del año {anio}, ante mí, ")
    _run(parrafo, notario["nombre"], negrita=True, subrayado=True)
    _run(parrafo, f", {notario['descripcion']}, comparecen: {apo['tratamiento']} ")
    _run(parrafo, apo["nombre"], negrita=True, subrayado=True)
    _run(parrafo,
        f", {apo['descripcion']}, quien comparece en representación, según se "
        "acreditará, del "
    )
    _run(parrafo, "BANCO DE CHILE", negrita=True)
    _run(parrafo,
        ", anónima de giro bancario, Rol Único Tributario número noventa y "
        "siete millones cuatro mil guion cinco, todos domiciliados en esta "
        "ciudad, calle Paseo Ahumada número doscientos cincuenta y uno, "
        "Santiago, ambos mayores de edad, quienes acreditan su identidad con "
        "la cédula respectiva y exponen: "
    )


def _personeria_banlegal(parrafo, apo: dict) -> None:
    """Personería y cierre para apoderado banlegal."""
    _run(parrafo, "LA PERSONERÍA", negrita=True, subrayado=True)
    _run(parrafo, f" de {apo['tratamiento']} ")
    _run(parrafo, f"{apo['nombre']}, ", negrita=True)
    _run(parrafo,
        f"para representar al Banco de Chile, consta de escritura pública de "
        f"fecha {apo['fecha_personeria']}, otorgada en {apo['notaria_personeria']}. "
        "Personería que no se inserta por ser conocidas de las partes y que el "
        "notario que autoriza ha tenido a la vista y vigente la citada personería. "
        "En comprobante firman, previa lectura. Se dio copia y se anotó en el "
        "LIBRO DE REPERTORIO con el número señalado. DOY FE."
    )


# ── Función pública ───────────────────────────────────────────────────────────

def formatear_banlegal(contenido: bytes, apo_id: str, notario: dict) -> bytes:
    """
    Formatea un borrador banlegal y retorna el .docx resultante en bytes.

    Args:
        contenido: bytes del .docx borrador banlegal.
        apo_id:    'ronnie' o 'felix'.
        notario:   dict del notario autorizante (de _NOTARIOS en formatear_controller).
    """
    if apo_id not in _APODERADOS_BANLEGAL:
        raise ValueError(f"Apoderado banlegal no registrado: '{apo_id}'")

    # Lazy imports para evitar dependencia circular con formatear_controller
    from .formatear_controller import (
        _configurar_pagina, _agregar_numero_pagina, _segmentar_cuerpo,
    )
    from .formatear_utils import (
        _es_documento_formateado, _extraer_datos_formateado, _extraer_clausulas_formateado,
    )

    try:
        doc_entrada = Document(io.BytesIO(contenido))
    except Exception as exc:
        logger.error("Error abriendo borrador banlegal: %s", exc)
        raise ValueError("El archivo no pudo ser leído. Verifica que sea un .docx válido.")

    apo  = _APODERADOS_BANLEGAL[apo_id]
    anio = __import__('num2words').num2words(datetime.now().year, lang='es')

    if _es_banlegal_formateado(doc_entrada):
        # Documento ya formateado por SAGE (cambio de compareciente) — re-extraer sin transformar
        datos          = _extraer_datos_formateado(doc_entrada)
        titulo         = datos.get("titulo", "")
        clausulas_tx   = _extraer_clausulas_formateado(doc_entrada)
        nombre_cliente = datos.get("nombre_cliente", "_________________________")
    else:
        # Borrador crudo — extracción y transformación normales.
        # Nota: _es_banlegal_formateado se usa en lugar de _es_documento_formateado
        # porque los borradores banlegal no tienen tabla interna (a diferencia de BDC estándar),
        # por lo que 'not doc.tables' siempre sería True y tomaría el camino equivocado.
        datos          = _extraer_datos_banlegal(doc_entrada)
        titulo         = _extraer_titulo_banlegal(doc_entrada)
        clausulas_raw  = _extraer_clausulas_banlegal(doc_entrada)
        nombre_cliente = _extraer_cliente_banlegal(clausulas_raw)
        clausulas_tx   = transformar_texto(clausulas_raw)

    # Detecta las comunas del Conservador para aplicar subrayado grueso (puede haber más de una)
    comunas = _extraer_comuna_conservador(doc_entrada)

    # ── Construcción del documento formateado ────────────────────────────────
    doc_salida = Document()
    _configurar_pagina(doc_salida, notario)
    _agregar_numero_pagina(doc_salida)

    for p in doc_salida.paragraphs:
        p._element.getparent().remove(p._element)

    JUST = WD_ALIGN_PARAGRAPH.JUSTIFY
    CTR  = WD_ALIGN_PARAGRAPH.CENTER
    LN2  = 2.0

    _parrafo(doc_salida, "J.E. /6 REPERTORIO Nº",                    negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, f"Abogado Redactor: {datos['abogado_redactor']}", negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, f"WF {datos['wf']}",                         negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, "FISCALÍA",                                  negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)

    p_titulo = _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _run(p_titulo, titulo, negrita=True, subrayado=True)
    _parrafo(doc_salida, "BANCO DE CHILE",  negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, "A",               negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, nombre_cliente,    negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)

    p_cuerpo = _parrafo(doc_salida, alineacion=JUST, interlineado=LN2)
    _comparecencia_banlegal(p_cuerpo, notario, apo, anio)

    segmentos = _segmentar_cuerpo(clausulas_tx, [nombre_cliente], comunas or ["BANLEGAL"])
    for seg in segmentos:
        _run(
            p_cuerpo,
            seg["texto"],
            negrita=seg["negrita"],
            subrayado=seg["subrayado"] == "DELGADO",
        )

    _personeria_banlegal(p_cuerpo, apo)

    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, apo["nombre"],    alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, "PP BANCO DE CHILE", alineacion=CTR, interlineado=LN2)

    buffer = io.BytesIO()
    doc_salida.save(buffer)
    buffer.seek(0)

    logger.info("Banlegal formateado — WF: %s | Apoderado: %s | Notario: %s",
                datos.get("wf"), apo_id, notario.get("header_nombre"))
    return buffer.read()
