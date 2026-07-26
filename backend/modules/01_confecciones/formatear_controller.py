"""
Módulo 01 — Confecciones de Borradores
Controller de formateo: transforma un .docx sin formato a escritura pública.

Proceso:
  1. Extrae datos del borrador (tabla interna o encabezado si ya fue formateado)
  2. Detecta comunas, nombres y co-deudores
  3. Construye el documento desde cero con formato notarial

Función pública:
  - formatear_documento(contenido, combo, notario_id) → bytes
"""

import io
import logging
import re
from datetime import datetime

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .formatear_clientes import procesar_nombres_borrador
from .formatear_utils import (
    transformar_texto,
    _extraer_datos_tabla,
    _extraer_titulo,
    _es_documento_formateado,
    _extraer_datos_formateado,
    _extraer_clausulas_formateado,
    _detectar_siguiente_clausula,
    _patron_nombre_flexible,
)

logger = logging.getLogger(__name__)

# ── Apoderados del Banco de Chile ─────────────────────────────────────────────
# Fuente: escritura pública de fecha 15 de junio de 2026, Notaría de Santiago
# de don Juan Francisco Alamos Ovejeros Notario Suplente.
# Deben firmar a lo menos 2 apoderados por escritura.
_APODERADOS: dict[str, dict] = {
    "nathalie": {
        "nombre":             "NATHALIE JEANNETTE VILLALOBOS RODRIGUEZ",
        "tratamiento":        "doña",
        "descripcion":        (
            "chilena, soltera, ingeniera en administración, cédula nacional de "
            "identidad número diecisiete millones ciento ochenta y dos mil "
            "novecientos seis guion dos"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
    "jonatan": {
        "nombre":             "JONATAN ENRIQUE GUTIERREZ AGUIRRE",
        "tratamiento":        "don",
        "descripcion":        (
            "chileno, casado, ingeniero en administración de empresas, cédula "
            "nacional de identidad número trece millones seiscientos sesenta y "
            "tres mil nueve guion uno"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
    "andrea": {
        "nombre":             "ANDREA MARITZA SANZ SAAVEDRA",
        "tratamiento":        "doña",
        "descripcion":        (
            "chilena, casada, contador auditor, cédula de identidad número once "
            "millones ochocientos ochenta y cinco mil cien guion cinco"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
    "carolina_v": {
        "nombre":             "CAROLINA DE LOS ANGELES VILCHES ESCOBAR",
        "tratamiento":        "doña",
        "descripcion":        (
            "chilena, divorciada, técnico en administración, cédula de identidad "
            "número trece millones ochocientos veintinueve mil trescientos "
            "cuarenta y dos guion cuatro"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
    "maria_ines": {
        "nombre":             "MARIA INES ASENJO ALARCON",
        "tratamiento":        "doña",
        "descripcion":        (
            "chilena, soltera, ingeniero comercial, cédula de identidad número "
            "nueve millones doscientos setenta y nueve mil setecientos sesenta "
            "y uno guion nueve"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
    "jacqueline": {
        "nombre":             "JACQUELINE DE LAS MERCEDES MIRANDA VALENZUELA",
        "tratamiento":        "doña",
        "descripcion":        (
            "chilena, divorciada, empleada bancaria, cédula de identidad número "
            "once millones ochocientos cincuenta y tres mil doscientos ochenta "
            "y uno guion tres"
        ),
        "fecha_personeria":   "quince de junio del año dos mil veintiséis",
        "notaria_personeria": "la Notaría de Santiago de don Juan Francisco Alamos Ovejeros Notario Suplente",
    },
}

# ── Notarios ──────────────────────────────────────────────────────────────────
# Las claves header_* aparecen en el encabezado de cada página impar impresa.
_NOTARIOS: dict[str, dict] = {
    "carolina": {
        "nombre":        "CAROLINA E. PIÑA CUEVAS",
        "descripcion":   (
            "chilena, divorciada, abogada, Notario Público Interino de la "
            "Trigésima Tercera Notaría de Santiago, cédula nacional de identidad "
            "número Dieciséis millones cuarenta y un mil seiscientos cinco guion K, "
            "con oficio en calle Huérfanos número novecientos setenta y nueve, "
            "oficina quinientos uno, de la comuna de Santiago"
        ),
        "header_nombre": "CAROLINA E. PIÑA CUEVAS",
        "header_cargo":  "NOTARIO PÚBLICO INTERINO",
        "header_oficio": "HUÉRFANOS 979 OF. 501 - SANTIAGO",
    },
    "suplente_1": {
        "nombre":        "BRENDA ANTONIA PÉREZ PINTO",
        "descripcion":   (
            "chilena, divorciada, abogado, cédula nacional de identidad número "
            "doce millones once mil cuatrocientos veintinueve guion ocho, "
            "Notario Público Suplente de la interina de la Trigésima Tercera "
            "Notaría de Santiago, doña CAROLINA E. PIÑA CUEVAS, según Decreto "
            "Judicial ya protocolizado, con oficio en calle Huérfanos número "
            "novecientos setenta y nueve, oficina quinientos uno de la Comuna "
            "de Santiago"
        ),
        "header_nombre": "BRENDA ANTONIA PÉREZ PINTO",
        "header_cargo":  "NOTARIO PÚBLICO SUPLENTE",
        "header_oficio": "HUÉRFANOS 979 OF. 501 - SANTIAGO",
    },
    # suplente_2: agregar cuando haya datos disponibles
}


# ── Helpers de formato ────────────────────────────────────────────────────────

def _set_lang_es(run) -> None:
    """Marca el run como español de Chile para que Word no lo subraye en rojo."""
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'es-CL')
    rPr.append(lang)


def _fuente(run: docx.text.run.Run) -> None:
    """Aplica fuente estándar de escritura pública al run."""
    run.font.name = "Courier New"
    run.font.size = Pt(12.5)
    _set_lang_es(run)


def _run(
    parrafo: docx.text.paragraph.Paragraph,
    texto: str,
    negrita: bool | None = None,
    subrayado: bool | None = None,
    subrayado_grueso: bool = False,
) -> None:
    """
    Agrega un run con formato notarial al párrafo.
    subrayado_grueso aplica THICK underline via XML (Conservador de Bienes Raíces).
    """
    if not texto:
        return
    r = parrafo.add_run(texto)
    _fuente(r)
    r.bold = negrita
    if subrayado_grueso:
        rPr = r._r.get_or_add_rPr()
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'thick')
        rPr.append(u)
        r.bold = True
    elif subrayado:
        r.underline = subrayado


def _parrafo(
    doc: Document,
    texto: str = "",
    negrita: bool | None = None,
    alineacion: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    interlineado: float = 2.0,
) -> docx.text.paragraph.Paragraph:
    """Agrega un párrafo simple (un run) al documento y lo retorna."""
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.line_spacing = interlineado
    p.paragraph_format.space_after  = Pt(0)
    if texto:
        _run(p, texto, negrita=negrita)
    return p


# ── Configuración de página y encabezado ─────────────────────────────────────

def _configurar_pagina(doc: Document, notario: dict) -> None:
    """
    Establece tamaño oficio, márgenes y encabezado con datos del notario.
    Encabezado solo en páginas impares via evenAndOddHeaders.
    """
    section = doc.sections[0]
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(33.02)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    configuracion = doc.settings.element
    configuracion.append(OxmlElement('w:evenAndOddHeaders'))

    def _linea_header(contenedor, txt: str, tamanio: int, negrita: bool) -> None:
        p = contenedor.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(txt)
        r.font.name = "Arial"
        r.font.size = Pt(tamanio)
        r.bold = negrita
        _set_lang_es(r)

    header_impar = section.header
    header_impar.is_linked_to_previous = False
    for p in header_impar.paragraphs:
        p._element.getparent().remove(p._element)
    _linea_header(header_impar, notario["header_nombre"], 13, True)
    _linea_header(header_impar, notario["header_cargo"],  13, True)
    _linea_header(header_impar, notario["header_oficio"], 10, False)
    _linea_header(header_impar, "",                       13, False)

    header_par = section.even_page_header
    header_par.is_linked_to_previous = False
    for p in header_par.paragraphs:
        p._element.getparent().remove(p._element)
    header_par.add_paragraph()


def _agregar_numero_pagina(doc: Document) -> None:
    """
    Inserta número de página centrado en el pie de todas las páginas.
    Con evenAndOddHeaders activo se deben configurar el pie impar y el pie par
    por separado; de lo contrario las páginas pares quedan sin número.
    """
    def _campo_page(contenedor) -> None:
        for p in contenedor.paragraphs:
            p._element.getparent().remove(p._element)
        p = contenedor.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        _fuente(r)
        inicio = OxmlElement('w:fldChar')
        inicio.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        fin = OxmlElement('w:fldChar')
        fin.set(qn('w:fldCharType'), 'end')
        r._r.append(inicio)
        r._r.append(instr)
        r._r.append(fin)

    section = doc.sections[0]
    _campo_page(section.footer)            # páginas impares
    _campo_page(section.even_page_footer)  # páginas pares


# ── Segmentación del cuerpo ───────────────────────────────────────────────────

def _segmentar_cuerpo(texto: str, nombres_clientes: list[str], comunas: list[str]) -> list[dict]:
    """
    Divide el texto de cláusulas en segmentos con su formato correspondiente.
    Cada segmento es {'texto', 'negrita', 'subrayado'} donde subrayado es
    None | 'DELGADO' | 'GRUESO'.
    Acepta múltiples comunas para documentos con más de un Conservador.
    """
    texto = re.sub(r'\bBanco de Chile\b', 'BANCO DE CHILE', texto, flags=re.IGNORECASE)
    # Normalizar cada mención del Conservador para que el patrón posterior funcione
    for comuna in comunas:
        texto = re.sub(
            r'Conservador\s+de\s+Bienes\s+Ra[ií]ces\s+de\s+' + re.escape(comuna),
            'Conservador de Bienes Raíces de ' + comuna, texto, flags=re.IGNORECASE,
        )

    # Patrón universal para encabezados de cláusula en español, incluyendo
    # ordinales compuestos hasta CENTESIMO (ej: CENTESIMO VIGESIMO CUARTO:)
    _DECENAS = (
        r'(?:CENTESIMO|NONAGESIMO|OCTAGESIMO|SEPTUAGESIMO|SEXAGESIMO|'
        r'QUINCUAGESIMO|CUADRAGESIMO|TRIGESIMO|VIGESIMO|DECIMO)'
    )
    _UNIDADES = r'(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|S[EÉ]PTIMO|OCTAVO|NOVENO)'
    _CLAUSULA_RE = (
        rf'(?:{_DECENAS}(?:\s+{_DECENAS})?(?:\s+{_UNIDADES})?|{_UNIDADES}):'
    )

    patrones = [
        (_CLAUSULA_RE, 'CLAUSULA', None),
        (r'BANCO DE CHILE',       True, None),
        (r'LA PERSON[EÉ]R[IÍ]A', True, 'DELGADO'),
    ]
    # Subrayado grueso para cada Conservador presente en el documento
    for comuna in comunas:
        patrones.append((r'Conservador de Bienes Ra[ií]ces de ' + re.escape(comuna), True, 'GRUESO'))
    for nombre in nombres_clientes:
        patrones.append((_patron_nombre_flexible(nombre), True, None))

    todos_matches: list[dict] = []
    for patron, negrita, subrayado in patrones:
        for m in re.finditer(patron, texto, re.UNICODE):
            todos_matches.append({
                "inicio": m.start(), "fin": m.end(),
                "texto": m.group(), "negrita": negrita, "subrayado": subrayado,
            })

    todos_matches.sort(key=lambda x: x["inicio"])
    sin_overlap: list[dict] = []
    ultimo_fin = 0
    for match in todos_matches:
        if match["inicio"] >= ultimo_fin:
            sin_overlap.append(match)
            ultimo_fin = match["fin"]

    segmentos: list[dict] = []
    pos = 0
    for match in sin_overlap:
        if pos < match["inicio"]:
            segmentos.append({"texto": texto[pos:match["inicio"]], "negrita": None, "subrayado": None})
        if match["negrita"] == "CLAUSULA":
            label = match["texto"][:-1]
            segmentos.append({"texto": label, "negrita": True, "subrayado": "DELGADO"})
            segmentos.append({"texto": ":",   "negrita": True, "subrayado": None})
        else:
            segmentos.append(match)
        pos = match["fin"]

    if pos < len(texto):
        segmentos.append({"texto": texto[pos:], "negrita": None, "subrayado": None})
    return segmentos


# ── Comparecencia y personería ────────────────────────────────────────────────

def _construir_comparecencia(parrafo, notario: dict, apo1: dict, apo2: dict, anio: str) -> None:
    """Agrega los runs de comparecencia notarial (dos apoderados BDC)."""
    _run(parrafo, "EN SANTIAGO DE CHILE, ", negrita=True)
    _run(parrafo, f"a ____________________ del año {anio}, ante mí, ")
    _run(parrafo, notario["nombre"], negrita=True, subrayado=True)
    _run(parrafo, f", {notario['descripcion']}, comparecen: {apo1['tratamiento']} ")
    _run(parrafo, apo1["nombre"], negrita=True, subrayado=True)
    _run(parrafo, f", {apo1['descripcion']}, y {apo2['tratamiento']} ")
    _run(parrafo, apo2["nombre"], negrita=True, subrayado=True)
    _run(parrafo,
        f", {apo2['descripcion']}, quienes comparecen en representación, "
        "según se acreditará, del "
    )
    _run(parrafo, "BANCO DE CHILE", negrita=True)
    _run(parrafo,
        ", anónima de giro bancario, Rol Único Tributario número noventa y "
        "siete millones cuatro mil guion cinco, todos domiciliados en esta "
        "ciudad, calle Paseo Ahumada número doscientos cincuenta y uno, "
        "Santiago, ambos mayores de edad, quienes acreditan su identidad con "
        "la cédula respectiva y exponen: "
    )


def _construir_personeria(parrafo, apo1: dict, apo2: dict) -> None:
    """Agrega personerías de los dos apoderados BDC y el cierre DOY FE."""
    for apo in (apo1, apo2):
        _run(parrafo, "LA PERSONERÍA", negrita=True, subrayado=True)
        _run(parrafo, f" de {apo['tratamiento']} ")
        _run(parrafo, f"{apo['nombre']}, ", negrita=True)
        _run(parrafo, "para representar al")
        _run(parrafo, " BANCO DE CHILE", negrita=True)
        _run(parrafo,
            f", consta de escritura pública de fecha {apo['fecha_personeria']}, "
            f"otorgada en {apo['notaria_personeria']}.- "
        )
    _run(parrafo,
        "La notario autorizante certifica que tuvo a la vista todas las "
        "personerías citadas en este instrumento y sus respectivas vigencias, "
        "y a ella le consta que tienen facultades para alzar y cancelar los "
        "gravámenes y prohibiciones. En comprobante firman, "
        "previa lectura. Se dio copia y se anotó en el LIBRO DE REPERTORIO "
        "con el número señalado. DOY FE."
    )


# ── Cláusula art. 401 N°12 COT ───────────────────────────────────────────────

_TEXTO_COT = (
    ": Atendido lo dispuesto en el artículo cuatrocientos uno número doce "
    "del Código Orgánico de Tribunales, el compareciente declara que se "
    "hará cargo de ingresar copia de la presente escritura pública al "
    "Conservador de Bienes Raíces respectivo, además de cubrir el costo "
    "de las inscripciones y anotaciones que correspondan, sin mediar "
    "intervención de la notario autorizante en la remisión de la "
    "referida copia. "
)


def _clausula_cot(parrafo, clausulas_tx: str) -> None:
    """
    Agrega la cláusula del art. 401 N°12 COT al párrafo de cuerpo.
    Detecta automáticamente el ordinal siguiente al último del cuerpo.
    El ordinal va en negrita + subrayado delgado; los dos puntos en negrita.
    Si la cláusula ya existe en el borrador, no se duplica.
    """
    if "cuatrocientos uno número doce" in clausulas_tx:
        return
    siguiente = _detectar_siguiente_clausula(clausulas_tx)
    _run(parrafo, f" {siguiente}", negrita=True, subrayado=True)
    _run(parrafo, _TEXTO_COT, negrita=None)


# ── Construcción del documento ────────────────────────────────────────────────

def _construir_documento(doc_salida: Document, datos: dict, titulo: str,
                          comunas: list[str], clausulas_tx: str,
                          nombres_en_cuerpo: list[str],
                          notario: dict, apo1: dict, apo2: dict,
                          anio: str) -> None:
    """Ensambla todos los bloques del documento formateado."""
    for p in doc_salida.paragraphs:
        p._element.getparent().remove(p._element)

    JUST = WD_ALIGN_PARAGRAPH.JUSTIFY
    CTR  = WD_ALIGN_PARAGRAPH.CENTER
    LN2  = 2.0

    # Encabezado: si hay múltiples conservadores se muestran todos separados por " / "
    cabecera_comuna = " / ".join(comunas) if comunas else "SIN COMUNA"

    _parrafo(doc_salida, "J.E. /6 REPERTORIO Nº",                         negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, f"Abogado Redactor: {datos['abogado_redactor']}", negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, f"WF {datos['wf']}",                              negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, cabecera_comuna,                                  negrita=True, alineacion=JUST, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)

    p_titulo = _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _run(p_titulo, titulo, negrita=True, subrayado=True)
    _parrafo(doc_salida, "BANCO DE CHILE",                                         negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, "A",                                                      negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, datos["nombre_cliente"] + datos.get("sufijo_titulo", ""), negrita=True, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)
    _parrafo(doc_salida, interlineado=LN2)

    p_cuerpo = _parrafo(doc_salida, alineacion=JUST, interlineado=LN2)
    _construir_comparecencia(p_cuerpo, notario, apo1, apo2, anio)

    _COMUNAS_ARIANNA = {"rancagua", "san miguel"}
    if any(c.lower() in _COMUNAS_ARIANNA for c in comunas):
        clausulas_tx = re.sub(
            r'Se faculta al portador de copia autorizada de la presente escritura,\s*'
            r'para requerir del Conservador de Bienes Ra[ií]ces respectivo,\s*'
            r'las anotaciones, inscripciones, subinscripciones y dem[aá]s actuaciones que procedan\.',
            'Se faculta a doña ARIANNA OSORIO, titular de la cédula de identidad número '
            'veintiséis millones ochenta y dos mil ochocientos once guion ocho '
            'para requerir del Conservador de Bienes Raíces respectivo, '
            'las anotaciones, inscripciones, subinscripciones y demás actuaciones que procedan.',
            clausulas_tx,
            flags=re.IGNORECASE,
        )
        nombres_en_cuerpo = list(nombres_en_cuerpo) + ["ARIANNA OSORIO"]

    for seg in _segmentar_cuerpo(clausulas_tx, nombres_en_cuerpo, comunas):
        _run(p_cuerpo, seg["texto"],
             negrita=seg["negrita"],
             subrayado=seg["subrayado"] == "DELGADO",
             subrayado_grueso=seg["subrayado"] == "GRUESO")

    # Cláusula art. 401 N°12 COT — ordinal detectado automáticamente
    _clausula_cot(p_cuerpo, clausulas_tx)

    _construir_personeria(p_cuerpo, apo1, apo2)

    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, apo1["nombre"],       alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, "PP. BANCO DE CHILE", alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, apo2["nombre"],       alineacion=CTR, interlineado=LN2)
    _parrafo(doc_salida, "PP. BANCO DE CHILE", alineacion=CTR, interlineado=LN2)


# ── Función pública principal ─────────────────────────────────────────────────

def formatear_documento(
    contenido: bytes,
    apo1_id: str,
    notario_id: str,
    apo2_id: str | None = None,
) -> tuple[bytes, str | None]:
    """
    Recibe el .docx sin formato y retorna (bytes_formateados, wf).
    wf es el número de carpeta extraído del documento, o None si no se pudo extraer.

    Si apo1_id es 'ronnie' o 'felix', delega a banlegal (apo2_id se ignora).
    Para BDC se requieren apo1_id y apo2_id distintos y registrados.
    """
    _KEYS_BANLEGAL = {'ronnie', 'felix'}
    if apo1_id in _KEYS_BANLEGAL:
        if notario_id not in _NOTARIOS:
            raise ValueError(f"Notario no registrado: '{notario_id}'")
        from .banlegal_controller import formatear_banlegal
        return formatear_banlegal(contenido, apo1_id, _NOTARIOS[notario_id]), None

    if apo1_id not in _APODERADOS:
        raise ValueError(f"Apoderado no registrado: '{apo1_id}'")
    if not apo2_id or apo2_id not in _APODERADOS:
        raise ValueError("Se requieren dos apoderados BDC válidos.")
    if notario_id not in _NOTARIOS:
        raise ValueError(f"Notario no registrado: '{notario_id}'")

    try:
        doc_entrada = Document(io.BytesIO(contenido))
    except Exception as exc:
        logger.error("Error abriendo el .docx: %s", exc)
        raise ValueError("El archivo no pudo ser leído. Verifica que sea un .docx válido.")

    # ── Extracción según tipo de documento ───────────────────────────────────
    if _es_documento_formateado(doc_entrada):
        datos        = _extraer_datos_formateado(doc_entrada)
        titulo       = datos.pop("titulo", "")
        comuna_raw   = datos.pop("comuna", "")
        clausulas_tx = _extraer_clausulas_formateado(doc_entrada)
        # Si no hay comuna en el encabezado (formato viejo), extraerla del cuerpo
        if not comuna_raw:
            from .formatear_utils import _extraer_comuna_conservador
            comunas_del_cuerpo = _extraer_comuna_conservador(doc_entrada)
            comunas = comunas_del_cuerpo if comunas_del_cuerpo else ["SIN COMUNA"]
        else:
            comunas = [c.strip() for c in comuna_raw.split("/") if c.strip()]
        datos["sufijo_titulo"] = ""
        # El título pudo haberse construido con sufijo " Y OTRO/A" concatenado al nombre.
        # Para buscar el nombre en el cuerpo hay que usar solo la parte base, sin sufijo.
        nombre_base = re.sub(r'\s+Y\s+OTR[AO]$', '', datos["nombre_cliente"], flags=re.IGNORECASE).strip()
        nombres_en_cuerpo = [nombre_base]
        # El encabezado del doc formateado guarda el nombre en orden notarial
        # (APELLIDO NOMBRE), pero el cuerpo lo tiene en orden original del banco
        # (NOMBRE APELLIDO). Agregar la forma de-invertida para que el resaltado
        # en negrita funcione independiente del orden en que aparezca.
        partes_nombre = nombre_base.split()
        if len(partes_nombre) > 2:
            forma_original = " ".join(partes_nombre[2:] + partes_nombre[:2])
            if forma_original not in nombres_en_cuerpo:
                nombres_en_cuerpo.append(forma_original)
        # Co-deudores en doc ya formateado: detectar "y don/doña NOMBRE" en cuerpo
        for m in re.finditer(
            r'[yY]\s+(?:don|do[n\u00f1]a)\s+'
            r'([A-Z\u00c0-\u00d6\u00d8-\u00de][A-Za-z\u00c0-\u00ff]+'
            r'(?:\s+[A-Z\u00c0-\u00d6\u00d8-\u00de][A-Za-z\u00c0-\u00ff]+){2,4})',
            clausulas_tx, re.UNICODE,
        ):
            nombre_codeudor = m.group(1).upper()
            if nombre_codeudor not in nombres_en_cuerpo:
                nombres_en_cuerpo.append(nombre_codeudor)
    else:
        datos  = _extraer_datos_tabla(doc_entrada)
        titulo = _extraer_titulo(doc_entrada)
        resultado = procesar_nombres_borrador(doc_entrada, datos)
        clausulas_tx      = resultado["clausulas_tx"]
        nombres_en_cuerpo = resultado["nombres_en_cuerpo"]
        comunas           = resultado["comuna"]

    notario     = _NOTARIOS[notario_id]
    apo1        = _APODERADOS[apo1_id]
    apo2        = _APODERADOS[apo2_id]
    anio_letras = __import__('num2words').num2words(datetime.now().year, lang='es')

    doc_salida = Document()
    _configurar_pagina(doc_salida, notario)
    _agregar_numero_pagina(doc_salida)
    _construir_documento(doc_salida, datos, titulo, comunas, clausulas_tx,
                         nombres_en_cuerpo, notario, apo1, apo2, anio_letras)

    buffer = io.BytesIO()
    doc_salida.save(buffer)
    buffer.seek(0)

    wf = datos.get("wf") or None
    logger.info("Documento formateado — WF: %s | Apo1: %s | Apo2: %s | Notario: %s", wf, apo1_id, apo2_id, notario_id)
    return buffer.read(), wf
