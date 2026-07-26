"""
Controlador de la tabla nombres_genero.

Responsabilidades:
  - Poblar la BD desde el archivo LISTAS NOMBRES.docx (seed inicial)
  - Mantener sets en memoria para búsquedas rápidas sin tocar la BD
  - Exponer obtener_tratamiento() para uso en formatear_controller

Funciones públicas:
  - poblar_desde_docx(ruta_docx, db) → int
  - inicializar_cache(db)            → None
  - obtener_tratamiento(primer_nombre) → "don" | "doña" | ""
"""

import logging
import unicodedata

from docx import Document
from sqlalchemy.orm import Session

from shared.nombres_model import NombreGenero

logger = logging.getLogger(__name__)

# Sets en memoria cargados al iniciar el servidor — evitan consultas a BD por request
_MASCULINOS: set[str] = set()
_FEMENINOS: set[str] = set()


# ── Utilidades internas ───────────────────────────────────────────────────────

def _normalizar(s: str) -> str:
    """Convierte a minúsculas y elimina tildes para comparación robusta."""
    return (
        unicodedata.normalize("NFD", s.lower())
        .encode("ascii", "ignore")
        .decode("ascii")
        .strip()
    )


# ── Seed desde docx ───────────────────────────────────────────────────────────

def poblar_desde_docx(ruta_docx: str, db: Session) -> int:
    """
    Lee un .docx de nombres e inserta los registros en la BD.
    Soporta dos formatos:
      - Tabla: 4 columnas (N°, NOMBRE HOMBRE, N°, NOMBRE MUJER) — nuevo formato
      - Párrafos: secciones con encabezados 'masculino' / 'femenino' — formato original
    Retorna la cantidad de nombres nuevos insertados (omite duplicados).
    """
    doc = Document(ruta_docx)
    insertados = 0

    ya_procesados: set[str] = {r.nombre for r in db.query(NombreGenero).all()}

    def _insertar(nombre_raw: str, genero: str) -> None:
        nonlocal insertados
        palabras = nombre_raw.strip().split()
        if not palabras:
            return
        clave = _normalizar(palabras[0])
        if len(clave) < 2 or clave in ya_procesados:
            return
        db.add(NombreGenero(nombre=clave, genero=genero))
        ya_procesados.add(clave)
        insertados += 1

    # Formato tabla: primera tabla con 4 columnas (N°, hombre, N°, mujer)
    tabla_nombres = next(
        (t for t in doc.tables if len(t.columns) == 4),
        None,
    )
    if tabla_nombres:
        for fila in tabla_nombres.rows[1:]:  # saltar encabezado
            hombre = fila.cells[1].text.strip()
            mujer  = fila.cells[3].text.strip()
            if hombre:
                _insertar(hombre, "M")
            if mujer:
                _insertar(mujer, "F")
    else:
        # Formato párrafos: secciones con encabezados masculino / femenino
        genero_actual: str | None = None
        for para in doc.paragraphs:
            texto = para.text.strip()
            if not texto:
                continue
            if "masculino" in texto.lower():
                genero_actual = "M"
                continue
            if "femenino" in texto.lower():
                genero_actual = "F"
                continue
            if genero_actual is None:
                continue
            for entrada in texto.split(","):
                _insertar(entrada, genero_actual)

    db.commit()
    logger.info("Seed de nombres completado: %d nuevos registros insertados.", insertados)
    return insertados


# ── Cache en memoria ──────────────────────────────────────────────────────────

def inicializar_cache(db: Session) -> None:
    """
    Carga todos los nombres desde BD a los sets en memoria.
    Debe llamarse UNA VEZ al arrancar el servidor (en main.py).
    """
    global _MASCULINOS, _FEMENINOS
    registros = db.query(NombreGenero).all()
    _MASCULINOS = {r.nombre for r in registros if r.genero == "M"}
    _FEMENINOS  = {r.nombre for r in registros if r.genero == "F"}
    logger.info(
        "Cache de nombres inicializado: %d masculinos, %d femeninos.",
        len(_MASCULINOS), len(_FEMENINOS),
    )


def seed_nombres_faltantes(db: Session) -> None:
    """
    Agrega nombres que se sabe que faltan en la BD.
    Llamado al arranque después de inicializar_cache para que el cache
    también se actualice con los nuevos registros.
    """
    faltantes = [
        ("cyntia", "F"),
    ]
    insertados = 0
    for clave, genero in faltantes:
        if not db.query(NombreGenero).filter(NombreGenero.nombre == clave).first():
            db.add(NombreGenero(nombre=clave, genero=genero))
            if genero == "M":
                _MASCULINOS.add(clave)
            else:
                _FEMENINOS.add(clave)
            insertados += 1
    if insertados:
        db.commit()
        logger.info("Seed de nombres faltantes: %d nuevos registros.", insertados)


def listar_nombres(db: Session) -> list[dict]:
    """Lista todos los nombres registrados ordenados alfabéticamente."""
    registros = db.query(NombreGenero).order_by(NombreGenero.nombre).all()
    return [{"nombre": r.nombre, "genero": r.genero} for r in registros]


def agregar_nombre(primer_nombre: str, genero: str, db: Session) -> dict:
    """
    Agrega un nombre a la BD y actualiza el cache en memoria.
    Si el nombre ya existe con distinto género, lo actualiza.
    """
    global _MASCULINOS, _FEMENINOS
    clave = _normalizar(primer_nombre)
    if len(clave) < 2:
        raise ValueError("El nombre es demasiado corto.")
    if genero not in ("M", "F"):
        raise ValueError("El género debe ser 'M' (masculino) o 'F' (femenino).")

    existente = db.query(NombreGenero).filter(NombreGenero.nombre == clave).first()
    if existente:
        if existente.genero == genero:
            raise ValueError(f"El nombre '{clave}' ya existe con ese género.")
        existente.genero = genero
        db.commit()
    else:
        db.add(NombreGenero(nombre=clave, genero=genero))
        db.commit()

    if genero == "M":
        _MASCULINOS.add(clave)
        _FEMENINOS.discard(clave)
    else:
        _FEMENINOS.add(clave)
        _MASCULINOS.discard(clave)

    return {"nombre": clave, "genero": genero}


def eliminar_nombre(primer_nombre: str, db: Session) -> None:
    """Elimina un nombre de la BD y del cache en memoria."""
    global _MASCULINOS, _FEMENINOS
    clave = _normalizar(primer_nombre)
    registro = db.query(NombreGenero).filter(NombreGenero.nombre == clave).first()
    if not registro:
        raise ValueError(f"Nombre '{clave}' no encontrado.")
    db.delete(registro)
    db.commit()
    _MASCULINOS.discard(clave)
    _FEMENINOS.discard(clave)


# ── Lookup de tratamiento ─────────────────────────────────────────────────────

def obtener_tratamiento(primer_nombre: str) -> str:
    """
    Retorna el tratamiento notarial según el primer nombre.
    Usa los sets en memoria — no requiere sesión de BD.

    Returns:
        "don"  si el nombre es masculino
        "doña" si el nombre es femenino
        ""     si el nombre no está registrado en la BD
    """
    clave = _normalizar(primer_nombre)
    if clave in _MASCULINOS:
        return "don"
    if clave in _FEMENINOS:
        return "doña"
    return ""
